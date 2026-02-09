"""
Generate AI Restyle Benchmark Newspost as DOCX with images
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_table(doc, headers, rows, header_color="003366"):
    """Add a styled table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], header_color)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = cell_text

    return table

def create_newspost():
    doc = Document()

    # Title
    title = doc.add_heading('Automated AI Restyle Benchmark Pipeline', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # TL;DR
    tldr = doc.add_paragraph()
    tldr_run = tldr.add_run('TL;DR: ')
    tldr_run.bold = True
    tldr.add_run('We built an automated evaluation system that replaces ad-hoc human reviews with a standardized, repeatable benchmark for AI Restyle quality. It uses Claude Code + Playwright + dual LLM judges (Gemini + Opus) to evaluate styles at scale.')
    tldr.paragraph_format.space_after = Pt(12)

    # Divider
    doc.add_paragraph('─' * 80)

    # Problem We Solved
    add_heading(doc, 'The Problem We Solved', 1)

    problem_headers = ['Before', 'After']
    problem_rows = [
        ['Evaluations were slow and didn\'t scale', 'Automated pipeline runs in minutes'],
        ['Results changed based on who was in the room', 'Consistent dual-LLM scoring every time'],
        ['No durable benchmark to compare versions', 'Persistent results in Excel for tracking'],
        ['Manual process made reuse hard', 'Reusable framework for any style/prompt'],
        ['Single evaluator bias', 'Cross-validation with Gemini + Opus 4.5'],
    ]
    add_table(doc, problem_headers, problem_rows)
    doc.add_paragraph()

    # How It Works
    add_heading(doc, 'How It Works', 1)

    flow = doc.add_paragraph()
    flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flow_text = '[Auto-discover styles] → [Apply to photos] → [ACRUE v3 Eval] → [Dual LLM scoring] → [Excel output]'
    flow_run = flow.add_run(flow_text)
    flow_run.font.name = 'Consolas'
    flow_run.font.size = Pt(11)
    flow_run.font.color.rgb = RGBColor(0, 120, 212)

    labels = doc.add_paragraph()
    labels.alignment = WD_ALIGN_PARAGRAPH.CENTER
    labels_text = '     Playwright            Automated           Assertions +        Gemini + Opus        Longitudinal'
    labels.add_run(labels_text).font.size = Pt(9)

    doc.add_paragraph()
    key_para = doc.add_paragraph()
    key_run = key_para.add_run('Key Innovation: ')
    key_run.bold = True
    key_para.add_run('The ACRUE v3 framework uses ')
    key_para.add_run('assertion-backed scoring').bold = True
    key_para.add_run(' - each quality dimension has Yes/No assertions with confidence levels (1-5) that must be answered before scoring. This grounds the evaluation and reduces hallucination.')

    # Demo Screenshots Section
    add_heading(doc, 'Demo: Playwright Automation in Action', 1)

    doc.add_paragraph('The automated pipeline captures the full user flow through OneDrive Photos AI Restyle:')

    # Add screenshots if they exist
    screenshots = [
        ('demo_01_gallery_view.png', 'Step 1: Gallery View - Navigate to OneDrive Photos'),
        ('demo_02_photo_viewer_restyle_option.png', 'Step 2: Photo Viewer - Open image with Restyle option visible'),
        ('demo_03_restyle_panel_styles.png', 'Step 3: Restyle Panel - 12 style presets available'),
        ('demo_04_ai_generating.png', 'Step 4: AI Generation - Processing the style transformation'),
    ]

    for img_file, caption in screenshots:
        if os.path.exists(img_file):
            doc.add_paragraph()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_file, width=Inches(5.5))

            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.italic = True
            cap.runs[0].font.size = Pt(10)
            cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # What You Can Do With It
    add_heading(doc, 'What You Can Do With It', 1)

    use_cases = [
        ('Compare prompt versions', '"Prompt A scores 18% higher than Prompt B on identity preservation"'),
        ('Benchmark new styles', 'Evaluate quality before shipping to users'),
        ('Track regressions', 'Automated comparison against baseline'),
        ('Separate feasibility from preference', 'Objective quality vs. subjective taste'),
        ('Scale evaluations', 'Run 100 image/style combinations overnight'),
    ]

    for i, (title, desc) in enumerate(use_cases, 1):
        para = doc.add_paragraph()
        para.add_run(f'{i}. {title}').bold = True
        para.add_run(f' - {desc}')

    # Artifacts Available
    add_heading(doc, 'Artifacts Available', 1)

    doc.add_heading('Presentations', level=2)
    pres_headers = ['File', 'Description']
    pres_rows = [
        ['Automated_Restyle_Benchmark_Slide.pptx', 'One-pager explaining the system'],
        ['Claude_Code_AI_Testing_LT.pptx', 'Leadership summary (13 slides)'],
        ['Claude_Code_Use_Cases_Presentation.pptx', 'Full use cases demo (21 slides)'],
    ]
    add_table(doc, pres_headers, pres_rows, "0078D4")
    doc.add_paragraph()

    doc.add_heading('Evaluation Framework', level=2)
    eval_headers = ['File', 'Description']
    eval_rows = [
        ['restyle_tests/acrue_v3_prompt.md', 'ACRUE v3 hybrid rubric (recommended)'],
        ['restyle_tests/run_acrue_v3.py', 'Python script for single/batch evals'],
        ['restyle_tests/style_assertions.json', 'Pre-defined assertions for 4 styles'],
        ['restyle_tests/batch_eval_config.json', 'Sample batch configuration'],
    ]
    add_table(doc, eval_headers, eval_rows, "107C10")
    doc.add_paragraph()

    doc.add_heading('Reports & Data', level=2)
    report_headers = ['File', 'Description']
    report_rows = [
        ['restyle_tests/results/acrue_v3_scores.json', 'Evaluation results output'],
        ['performance_metrics.csv', 'Latency data (28 data points)'],
        ['AI_Restyle_Bug_Bash_Report.md', 'Bug discovery summary'],
    ]
    add_table(doc, report_headers, report_rows, "D83B01")
    doc.add_paragraph()

    # ACRUE v3 Scoring
    add_heading(doc, 'ACRUE v3 Scoring Dimensions', 1)

    acrue_headers = ['Dimension', 'Weight', 'What It Measures']
    acrue_rows = [
        ['A - Accuracy', '1.0', 'Style transformation quality'],
        ['C - Completeness', '1.0', 'Full image coverage'],
        ['R - Relevance', '0.5', 'Appropriate for restyle context'],
        ['U - Usefulness', '0.5', 'User would want to save/share'],
        ['E - Exceptional', '2.0', 'Wow factor, creativity'],
    ]
    add_table(doc, acrue_headers, acrue_rows)
    doc.add_paragraph()

    grade_para = doc.add_paragraph()
    grade_para.add_run('Grade Thresholds: ').bold = True
    grade_para.add_run('90%+ = A+, 80%+ = A, 70%+ = B, 60%+ = C, <60% = F')

    # Quick Start
    add_heading(doc, 'Quick Start', 1)

    code_block = """# Single evaluation
python restyle_tests/run_acrue_v3.py -o photo.jpg -r styled.png -s "Anime"

# Batch evaluation
python restyle_tests/run_acrue_v3.py --batch restyle_tests/batch_eval_config.json

# Plan-only mode (preview assertions without running)
python restyle_tests/run_acrue_v3.py -o photo.jpg -r styled.png -s "Storybook" --plan-only"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code_block)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_paragraph()
    prereq = doc.add_paragraph()
    prereq.add_run('Prerequisites: ').bold = True
    prereq.add_run('pip install google-generativeai python-docx | Set GEMINI_API_KEY environment variable')

    # What's Next
    add_heading(doc, "What's Next", 1)

    roadmap = [
        'CI/CD integration for automated quality gates',
        'Parallel agent testing for faster throughput',
        'Regression comparison across builds',
        'Multi-browser testing (Edge, Chrome, Safari)',
        'Expand style assertions library',
    ]

    for item in roadmap:
        para = doc.add_paragraph(item, style='List Bullet')

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 80)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Tech Stack: ').bold = True
    footer.add_run('Claude Code | Playwright MCP | Gemini 2.0 Flash | Opus 4.5 | ACRUE v3 Framework')

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.add_run('Built with Claude Code - autonomous AI testing for OneDrive Photos').italic = True

    # Save
    output_path = 'AI_Restyle_Benchmark_Newspost.docx'
    doc.save(output_path)
    print(f'Created: {output_path}')
    return output_path

if __name__ == '__main__':
    create_newspost()
