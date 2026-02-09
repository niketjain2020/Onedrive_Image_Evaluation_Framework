"""
Generate: Why We Built This - Embracing Flawed Systems Story
A comparison of Human Evals vs Multi-LLM Orchestrated Evals
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_quote_block(doc, text):
    """Add an indented quote block"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.right_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(f'"{text}"')
    run.font.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 90, 158)
    return para

def add_section_break(doc):
    doc.add_paragraph()
    para = doc.add_paragraph('─' * 60)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].font.color.rgb = RGBColor(200, 200, 200)
    doc.add_paragraph()

def create_story_doc():
    doc = Document()

    # ===== TITLE =====
    title = doc.add_heading('Why We Built This', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('Embracing Flawed Systems to Get Better Outcomes')
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(0, 120, 212)
    sub_run.font.italic = True

    doc.add_paragraph()

    # ===== OPENING =====
    doc.add_heading('The Starting Point', 1)

    p = doc.add_paragraph()
    p.add_run('This work did not start from the idea of building a perfect evaluation framework.').bold = True

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.add_run('It started from the opposite realization: ').font.size = Pt(11)
    p2.add_run('every component we use is flawed.').bold = True

    doc.add_paragraph()

    # Flaws list - styled as impactful statements
    flaws = [
        'LLMs hallucinate.',
        'Browser automation is brittle.',
        'Evaluation rubrics are imperfect.',
        'Human judgment is inconsistent and slow.',
    ]
    for flaw in flaws:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(flaw)
        run.font.color.rgb = RGBColor(150, 0, 0)
        run.font.size = Pt(12)

    doc.add_paragraph()

    # Key insight callout
    insight = doc.add_paragraph()
    insight.paragraph_format.left_indent = Inches(0.3)
    insight.add_run('→ ').font.color.rgb = RGBColor(0, 120, 212)
    insight.add_run('Instead of trying to eliminate these flaws, we built a system that ').font.size = Pt(11)
    insight.add_run('composes them deliberately.').bold = True

    add_section_break(doc)

    # ===== CORE INSIGHT =====
    doc.add_heading('The Core Insight: Orchestration Beats Perfection', 1)

    doc.add_paragraph('The real value of this system is not any single model, metric, or assertion.')

    doc.add_paragraph()
    highlight = doc.add_paragraph()
    hl_run = highlight.add_run('The value comes from orchestrating multiple imperfect systems and letting them reason over the same inputs in different ways.')
    hl_run.bold = True
    hl_run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph('We use:')

    components = [
        ('Playwright MCP', 'to interact with real product surfaces'),
        ('Multiple LLM MCPs', 'with different strengths'),
        ('A structured pipeline', 'that forces consistency across runs'),
    ]
    for name, desc in components:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        para.add_run(f'• {name} ').bold = True
        para.add_run(desc)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Each component is flawed on its own.').italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Together, they create a stable, repeatable signal.').bold = True

    add_quote_block(doc, "This is not about trusting one model. It is about triangulating across models.")

    add_section_break(doc)

    # ===== WHY MULTIPLE JUDGES =====
    doc.add_heading('Why Multiple LLM Judges Matter', 1)

    p = doc.add_paragraph()
    p.add_run('Different models fail differently.').bold = True

    doc.add_paragraph()

    model_traits = [
        'Some models are more technical and spec-driven',
        'Others have a stronger creative or aesthetic bias',
        'Some are better at recall and coverage',
        'Others are better at holistic judgment',
    ]
    for trait in model_traits:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        para.add_run(f'• {trait}')

    doc.add_paragraph()
    doc.add_paragraph('By running the same inputs through multiple judges, we:')
    doc.add_paragraph()

    benefits = [
        ('Expose disagreement explicitly', 'No more hidden consensus bias'),
        ('Reduce blind spots', 'from any single model'),
        ('Get a stronger starting signal', 'before human review'),
    ]
    for i, (benefit, detail) in enumerate(benefits, 1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        para.add_run(f'{i}. {benefit}').bold = True
        para.add_run(f' — {detail}')

    add_quote_block(doc, 'This is fundamentally different from asking one model, or one group of humans, "what do you think?"')

    add_section_break(doc)

    # ===== HOW THIS COMPARES =====
    doc.add_heading('How This Compares to Human Evaluation', 1)

    doc.add_paragraph('Historically, we relied on human evals:')
    doc.add_paragraph()

    human_problems = [
        'They are slow and expensive',
        'Results change depending on who is reviewing',
        'Benchmarks are hard to preserve over time',
        'Past decisions are difficult to revisit or compare against',
    ]
    for problem in human_problems:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        run = para.add_run(f'• {problem}')
        run.font.color.rgb = RGBColor(150, 0, 0)

    doc.add_paragraph()

    key = doc.add_paragraph()
    key.add_run('Most importantly, human eval does not compound.').bold = True

    doc.add_paragraph('Every new review starts from scratch.')

    doc.add_paragraph()

    contrast = doc.add_paragraph()
    contrast.add_run('This system does.').bold = True
    contrast.runs[0].font.size = Pt(14)
    contrast.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph()
    doc.add_paragraph('Each run adds:')

    adds = ['More context', 'More comparable data', 'More historical signal']
    for item in adds:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        para.add_run(f'✓ {item}').font.color.rgb = RGBColor(0, 128, 0)

    add_quote_block(doc, "The system itself does not become perfect, but it becomes more informative.")

    # ===== COMPARISON TABLE =====
    doc.add_paragraph()
    doc.add_heading('Side-by-Side Comparison', 2)

    # Add image if exists
    if os.path.exists('demo_03_restyle_panel_styles.png'):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture('demo_03_restyle_panel_styles.png', width=Inches(5))
        cap = doc.add_paragraph('AI Restyle: The first application of this evaluation system')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.italic = True
        cap.runs[0].font.size = Pt(10)
        cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph()

    # Comparison table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    headers = ['Dimension', 'Human Evaluation', 'Multi-LLM Orchestration']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], "003366")

    comparisons = [
        ['Speed', 'Hours to days per batch', 'Minutes per batch'],
        ['Consistency', 'Varies by reviewer', 'Same rubric every time'],
        ['Scalability', 'Expensive to scale', '100+ evals overnight'],
        ['Bias', 'Single perspective', 'Multiple models triangulate'],
        ['Memory', 'Decisions forgotten', 'Every run logged & comparable'],
        ['Compounding', 'Starts from scratch', 'Adds historical signal'],
        ['Disagreement', 'Hidden in group dynamics', 'Exposed explicitly'],
        ['Cost', 'High (human time)', 'Low (API calls)'],
    ]

    for row_data in comparisons:
        row = table.add_row()
        row.cells[0].text = row_data[0]
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = row_data[1]
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(150, 0, 0)
        row.cells[2].text = row_data[2]
        row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

    add_section_break(doc)

    # ===== GENERALIZATION =====
    doc.add_heading('Why This Generalizes Beyond Images', 1)

    doc.add_paragraph('Image restyle is just the first application.')

    doc.add_paragraph()
    doc.add_paragraph('The same approach works anywhere we have:')

    prereqs = ['Clear inputs', 'Observable outputs', 'A rubric we care about']
    for p in prereqs:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        para.add_run(f'✓ {p}').font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph()
    doc.add_paragraph('That includes:')

    applications = [
        ('Prompt optimization', 'Which prompt generates better outputs?'),
        ('Recall and ranking quality', 'Are we surfacing the right results?'),
        ('Retrieval and recommendation', 'Does the system find relevant content?'),
        ('Any AI system where judgment matters', 'Anywhere quality is subjective'),
    ]
    for app, desc in applications:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        para.add_run(f'• {app}').bold = True
        para.add_run(f' — {desc}')

    add_quote_block(doc, "The hard part is defining the pipeline upfront. Once that is done, the orchestration scales.")

    add_section_break(doc)

    # ===== THE REAL OUTCOME =====
    doc.add_heading('The Real Outcome', 1)

    p = doc.add_paragraph()
    p.add_run('This is not about replacing humans.').font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('It is about giving humans a better starting point.').bold = True
    p.runs[0].font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph('Instead of debating from scratch, we start from:')

    outcomes = ['Structured disagreement', 'Multi-model reasoning', 'Historical benchmarks']
    for outcome in outcomes:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.3)
        para.add_run(f'→ {outcome}').font.color.rgb = RGBColor(0, 90, 158)
        para.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # Final statement - make it prominent
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.add_run('We move faster not because the system is perfect,').italic = True

    final2 = doc.add_paragraph()
    final2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final2.add_run('but because it is consistently imperfect in the same way.').bold = True
    final2.runs[0].font.size = Pt(14)

    doc.add_paragraph()

    # Closing quote
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(20)
    run = closing.add_run('And that turns subjective judgment into something we can actually build on over time.')
    run.font.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 90, 158)

    # ===== DEMO SCREENSHOTS =====
    add_section_break(doc)
    doc.add_heading('The System in Action', 1)

    doc.add_paragraph('Playwright MCP automates the full user flow:')

    screenshots = [
        ('demo_01_gallery_view.png', 'Gallery View - Automated navigation to OneDrive Photos'),
        ('demo_04_ai_generating.png', 'AI Generation - Capturing the restyle process'),
    ]

    for img_file, caption in screenshots:
        if os.path.exists(img_file):
            doc.add_paragraph()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_file, width=Inches(5.5))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.italic = True
            cap.runs[0].font.size = Pt(10)
            cap.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    # ===== FOOTER =====
    doc.add_paragraph()
    doc.add_paragraph('─' * 80)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Tech Stack: ').bold = True
    footer.add_run('Claude Code | Playwright MCP | Gemini 2.0 Flash | Opus 4.5 | ACRUE v3')

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.add_run('Multi-LLM orchestration for AI evaluation at scale').italic = True

    # Save
    output_path = 'AI_Restyle_Benchmark_Story.docx'
    doc.save(output_path)
    print(f'Created: {output_path}')
    return output_path

if __name__ == '__main__':
    create_story_doc()
