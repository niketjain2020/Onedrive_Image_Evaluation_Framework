"""
Generate ACRUE Evaluation Report in DOCX format with Before/After Images
"""
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# Image path mappings for before/after comparisons
IMAGE_MAPPINGS = {
    "Renaissance Oil Painting": {
        "original": "test1_adult_original.png",
        "restyled": "prompt01_renaissance_oil_painting.png"
    },
    "Japanese Ukiyo-e": {
        "original": "test1_adult_original.png",
        "restyled": "prompt05_japanese_ukiyoe.png"
    },
    "Movie Poster": {
        "original": "test1_adult_original.png",
        "restyled": "img1-01-movie-poster.png"
    },
    "Anime": {
        "original": "test1_adult_original.png",
        "restyled": "img1-03-anime.png"
    }
}


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def get_grade_color(grade):
    """Return color based on grade."""
    colors = {
        "A+": "90EE90",  # Light green
        "A": "98FB98",   # Pale green
        "B": "FFFFE0",   # Light yellow
        "C": "FFD700",   # Gold
        "F": "FF6B6B"    # Light red
    }
    return colors.get(grade, "FFFFFF")


def create_acrue_report(json_path, output_path=None):
    """Create DOCX report from ACRUE evaluation results."""

    # Load results
    with open(json_path, "r", encoding="utf-8") as f:
        results = json.load(f)

    if not isinstance(results, list):
        results = [results]

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading("ACRUE Image Evaluation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)

    summary_table = doc.add_table(rows=1, cols=5)
    summary_table.style = "Table Grid"
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["#", "Style", "Grade", "Score", "Percentage"]
    header_cells = summary_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], "4472C4")
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for idx, result in enumerate(results, 1):
        row = summary_table.add_row()
        summary = result.get("summary", {})
        grade = summary.get("grade", "N/A")

        row.cells[0].text = str(idx)
        row.cells[1].text = result.get("style", "Unknown")
        row.cells[2].text = grade
        row.cells[3].text = f"{summary.get('weighted_total', 0):.2f}/25"
        row.cells[4].text = f"{summary.get('percentage', 0):.1f}%"

        # Color grade cell
        set_cell_shading(row.cells[2], get_grade_color(grade))

    doc.add_paragraph()

    # ACRUE Framework Overview
    doc.add_heading("ACRUE Framework (Restyle-Optimized)", level=1)

    intro_para = doc.add_paragraph()
    intro_para.add_run("Key Insight: ").bold = True
    intro_para.add_run("AI Restyle is ")
    intro_para.add_run("style transfer").bold = True
    intro_para.add_run(", not content editing. The rubric focuses on identity preservation and style authenticity.")

    doc.add_paragraph()

    # What stays vs what changes table
    change_table = doc.add_table(rows=6, cols=2)
    change_table.style = "Table Grid"

    change_table.rows[0].cells[0].text = "What STAYS the Same"
    change_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    set_cell_shading(change_table.rows[0].cells[0], "90EE90")

    change_table.rows[0].cells[1].text = "What CHANGES"
    change_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    set_cell_shading(change_table.rows[0].cells[1], "FFB347")

    changes = [
        ("Subject identity (recognizable)", "Visual art style"),
        ("Composition & layout", "Color palette & lighting"),
        ("Pose & positioning", "Texture & brushwork"),
        ("Scene content", "Mood & atmosphere"),
        ("Object relationships", "Medium (photo → art)")
    ]
    for i, (stays, changes_to) in enumerate(changes, 1):
        change_table.rows[i].cells[0].text = stays
        change_table.rows[i].cells[1].text = changes_to

    doc.add_paragraph()

    framework_text = """The ACRUE framework evaluates AI-restyled images across five weighted dimensions:

• Accuracy (Weight: 1.0) - Identity preservation, style authenticity, composition fidelity, appropriate additions
• Completeness (Weight: 1.0) - Subject coverage, background harmony, style saturation, global consistency, structural integrity
• Relevance (Weight: 0.5) - Style match, mood alignment
• Usefulness (Weight: 0.5) - Share-ready quality, visual clarity, artifact-free, practical quality
• Exceptional Value (Weight: 2.0) - Style mastery, creative enhancement, emotional impact, distinctive character, shareability

Grade Thresholds: A+ (90-100%), A (80-89%), B (70-79%), C (60-69%), F (<60%)"""

    doc.add_paragraph(framework_text)
    doc.add_paragraph()

    # Detailed Results
    doc.add_heading("Detailed Evaluation Results", level=1)

    # Base path for images
    base_img_path = Path(__file__).resolve().parent.parent.parent / ".playwright-mcp"

    for idx, result in enumerate(results, 1):
        # Evaluation header
        doc.add_heading(f"Evaluation {idx}: {result.get('style', 'Unknown')}", level=2)

        # Before/After Images
        style_name = result.get("style", "")
        img_mapping = IMAGE_MAPPINGS.get(style_name, {})

        if img_mapping:
            doc.add_heading("Before & After Comparison", level=3)

            # Create 2x2 table for images and labels
            img_table = doc.add_table(rows=2, cols=2)
            img_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            original_path = base_img_path / img_mapping.get("original", "")
            restyled_path = base_img_path / img_mapping.get("restyled", "")

            # Add original image
            if original_path.exists():
                try:
                    img_table.rows[0].cells[0].paragraphs[0].add_run().add_picture(
                        str(original_path), width=Inches(2.5)
                    )
                except Exception as e:
                    img_table.rows[0].cells[0].text = f"[Original: {original_path.name}]"
            else:
                img_table.rows[0].cells[0].text = "[Original Image Not Found]"

            # Add restyled image
            if restyled_path.exists():
                try:
                    img_table.rows[0].cells[1].paragraphs[0].add_run().add_picture(
                        str(restyled_path), width=Inches(2.5)
                    )
                except Exception as e:
                    img_table.rows[0].cells[1].text = f"[Restyled: {restyled_path.name}]"
            else:
                img_table.rows[0].cells[1].text = "[Restyled Image Not Found]"

            # Labels
            img_table.rows[1].cells[0].text = "ORIGINAL"
            img_table.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_table.rows[1].cells[0].paragraphs[0].runs[0].bold = True

            img_table.rows[1].cells[1].text = style_name.upper()
            img_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_table.rows[1].cells[1].paragraphs[0].runs[0].bold = True

            doc.add_paragraph()

        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run("Evaluation ID: ").bold = True
        meta_para.add_run(result.get("evaluation_id", "N/A") + "\n")
        meta_para.add_run("Timestamp: ").bold = True
        meta_para.add_run(result.get("timestamp", "N/A") + "\n")
        meta_para.add_run("Rubric Version: ").bold = True
        meta_para.add_run(result.get("rubric_version", "unknown") + "\n")
        meta_para.add_run("Original Image: ").bold = True
        meta_para.add_run(Path(result.get("original_image", "N/A")).name + "\n")
        meta_para.add_run("Restyled Image: ").bold = True
        meta_para.add_run(Path(result.get("restyled_image", "N/A")).name)

        # Summary box
        summary = result.get("summary", {})
        grade = summary.get("grade", "N/A")

        grade_para = doc.add_paragraph()
        grade_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        grade_run = grade_para.add_run(f"GRADE: {grade}  |  SCORE: {summary.get('weighted_total', 0):.2f}/25  |  {summary.get('percentage', 0):.1f}%")
        grade_run.bold = True
        grade_run.font.size = Pt(14)

        # Overall Assessment
        assessment = result.get("overall_assessment", "")
        if assessment:
            doc.add_paragraph()
            assess_para = doc.add_paragraph()
            assess_para.add_run("Overall Assessment: ").bold = True
            assess_para.add_run(assessment)

        doc.add_paragraph()

        # Dimension scores table
        doc.add_heading("Dimension Scores", level=3)

        scores = result.get("scores", {})
        dim_table = doc.add_table(rows=1, cols=4)
        dim_table.style = "Table Grid"

        # Header
        dim_headers = ["Dimension", "Score", "Weight", "Weighted Score"]
        for i, h in enumerate(dim_headers):
            dim_table.rows[0].cells[i].text = h
            dim_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            set_cell_shading(dim_table.rows[0].cells[i], "D9E2F3")

        # Dimension rows
        dimension_names = {
            "accuracy": "Accuracy",
            "completeness": "Completeness",
            "relevance": "Relevance",
            "usefulness": "Usefulness",
            "exceptional_value": "Exceptional Value"
        }

        for dim_key, dim_name in dimension_names.items():
            dim_data = scores.get(dim_key, {})
            row = dim_table.add_row()
            row.cells[0].text = dim_name
            row.cells[1].text = f"{dim_data.get('dimension_score', 0):.1f}/5.0"
            row.cells[2].text = f"×{dim_data.get('weight', 0)}"
            row.cells[3].text = f"{dim_data.get('weighted_score', 0):.2f}"

        # Total row
        total_row = dim_table.add_row()
        total_row.cells[0].text = "TOTAL"
        total_row.cells[0].paragraphs[0].runs[0].bold = True
        total_row.cells[1].text = ""
        total_row.cells[2].text = ""
        total_row.cells[3].text = f"{summary.get('weighted_total', 0):.2f}/25"
        total_row.cells[3].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Sub-scores detail
        doc.add_heading("Sub-Score Details", level=3)

        for dim_key, dim_name in dimension_names.items():
            dim_data = scores.get(dim_key, {})
            sub_scores = dim_data.get("sub_scores", {})

            if sub_scores:
                doc.add_paragraph(f"{dim_name}:", style="List Bullet")

                for sub_key, sub_data in sub_scores.items():
                    sub_name = sub_key.replace("_", " ").title()
                    score = sub_data.get("score", 0)
                    rationale = sub_data.get("rationale", "")

                    sub_para = doc.add_paragraph(style="List Bullet 2")
                    sub_para.add_run(f"{sub_name}: {score}/5").bold = True
                    if rationale:
                        sub_para.add_run(f" - {rationale}")

        # Page break between evaluations (except last)
        if idx < len(results):
            doc.add_page_break()

    # Save document
    if output_path is None:
        output_path = Path(__file__).resolve().parent.parent / "results" / "ACRUE_Evaluation_Report.docx"

    output_path = Path(output_path)
    output_path.parent.mkdir(exist_ok=True)

    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    json_path = Path(__file__).resolve().parent.parent / "results" / "acrue_scores.json"
    create_acrue_report(json_path)
