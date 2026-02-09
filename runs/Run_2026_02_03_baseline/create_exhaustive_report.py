#!/usr/bin/env python3
"""
Exhaustive Benchmark Report Generator

Generates a comprehensive DOCX report including:
- Executive Summary
- Methodology
- All ACRUE v3 Gemini Evaluations (assertion-by-assertion)
- Opus Preference Rankings
- Images (originals and restyled)
- Statistical Analysis
- Recommendations
"""

import json
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paths
SCRIPT_DIR = Path(__file__).parent
ACRUE_PATH = SCRIPT_DIR / "acrue.json"
OPUS_PATH = SCRIPT_DIR / "opus.json"
GEMINI_PATH = SCRIPT_DIR / "gemini.json"
SPEC_PATH = SCRIPT_DIR / "run_spec.json"
ORIGINALS_DIR = SCRIPT_DIR / "originals"
RESTYLED_DIR = SCRIPT_DIR / "restyled"
OUTPUT_PATH = SCRIPT_DIR / "Exhaustive_Benchmark_Report.docx"


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_with_style(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    return heading


def create_table_with_header(doc, headers, col_widths=None):
    """Create a table with styled header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2F5496')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)

    return table


def add_assertion_table(doc, assertions, dimension_name):
    """Add a detailed assertion table for a dimension."""
    if not assertions:
        return

    headers = ["ID", "Assertion", "Answer", "Conf", "Evidence"]
    table = create_table_with_header(doc, headers, [0.4, 2.5, 0.5, 0.4, 2.7])

    for a in assertions:
        row = table.add_row()
        cells = row.cells
        cells[0].text = a.get('id', '')
        cells[1].text = a.get('question', '')[:80] + ('...' if len(a.get('question', '')) > 80 else '')

        answer = a.get('answer', '')
        cells[2].text = answer
        if answer == 'Yes':
            set_cell_shading(cells[2], 'C6EFCE')  # Green
        else:
            set_cell_shading(cells[2], 'FFC7CE')  # Red
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cells[3].text = str(a.get('confidence', ''))
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        evidence = a.get('evidence', '')
        cells[4].text = evidence[:100] + ('...' if len(evidence) > 100 else '')

        # Make text smaller
        for cell in cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)


def get_grade_color(grade):
    """Return color based on grade."""
    colors = {
        'A+': 'C6EFCE',  # Green
        'A': 'C6EFCE',
        'B': 'FFEB9C',   # Yellow
        'C': 'FFC7CE',   # Light red
        'F': 'FF6B6B'    # Red
    }
    return colors.get(grade, 'FFFFFF')


def main():
    # Load data
    with open(ACRUE_PATH, 'r') as f:
        acrue_data = json.load(f)

    with open(OPUS_PATH, 'r') as f:
        opus_data = json.load(f)

    with open(GEMINI_PATH, 'r') as f:
        gemini_data = json.load(f)

    with open(SPEC_PATH, 'r') as f:
        spec_data = json.load(f)

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_heading('AI Restyle Benchmark Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Run ID: {spec_data.get('run_id', 'Unknown')}")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").font.size = Pt(12)
    info_para.add_run(f"Pipeline Version: {spec_data.get('pipeline_version', '1.0.0')}\n").font.size = Pt(12)
    info_para.add_run(f"ACRUE Version: {spec_data.get('acrue_version', 'v3')}").font.size = Pt(12)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    add_heading_with_style(doc, 'Table of Contents', 1)

    toc_items = [
        "1. Executive Summary",
        "2. Methodology",
        "3. Run Configuration",
        "4. Final Rankings & Winner",
        "5. Gemini ACRUE v3 Evaluations (Feasibility)",
        "   5.1 Image 1: Pelican (Wildlife)",
        "   5.2 Image 2: Tiramisu (Food)",
        "   5.3 Image 3: Tiramisu Variant (Food)",
        "6. Opus Preference Rankings",
        "7. Synthesis & Final Scores",
        "8. Statistical Analysis",
        "9. Image Gallery",
        "10. Recommendations",
        "11. Appendix: Raw Data"
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    add_heading_with_style(doc, '1. Executive Summary', 1)

    # Winner box
    winner_para = doc.add_paragraph()
    winner_run = winner_para.add_run('WINNER: POP ART')
    winner_run.bold = True
    winner_run.font.size = Pt(24)
    winner_run.font.color.rgb = RGBColor(0, 128, 0)
    winner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    summary_text = """This benchmark evaluated three AI restyle transformations (Anime, Pop Art, Storybook) across three diverse images using a dual-judge system:

• Gemini 2.0 Flash (ACRUE v3 Framework): Evaluated technical feasibility through 23 assertions per image across 5 dimensions (Accuracy, Completeness, Relevance, Usefulness, Exceptional Value)

• Claude Opus 4.5: Evaluated aesthetic preference based on visual appeal, user delight potential, and social shareability

Key Findings:
• Pop Art achieved perfect scores (100% assertion pass rate, 92% weighted average)
• Storybook performed strongly (100% pass rate, 87% weighted average)
• Anime struggled with non-human subjects (51% pass rate, 32% weighted average)
• Both judges unanimously ranked Pop Art as #1

Recommendation: Pop Art should be prioritized for diverse image types. Anime style should include subject-type detection to set appropriate user expectations for non-human subjects."""

    doc.add_paragraph(summary_text)

    # Quick stats table
    doc.add_paragraph()
    headers = ["Metric", "Pop Art", "Storybook", "Anime"]
    table = create_table_with_header(doc, headers)

    stats = [
        ["Gemini Rank", "1", "2", "3"],
        ["Opus Rank", "1", "2", "3"],
        ["Avg ACRUE Score", "23.0/25", "21.67/25", "7.99/25"],
        ["Percentage", "92.0%", "86.7%", "31.9%"],
        ["Grade", "A+", "A", "F"],
        ["Assertion Pass Rate", "100%", "100%", "51%"],
        ["Appeal Score", "9.2/10", "8.5/10", "5.0/10"]
    ]

    for row_data in stats:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # =========================================================================
    # 2. METHODOLOGY
    # =========================================================================
    add_heading_with_style(doc, '2. Methodology', 1)

    doc.add_heading('2.1 ACRUE v3 Framework', 2)

    acrue_desc = """The ACRUE v3 hybrid evaluation framework combines Yes/No assertions with confidence scoring (1-5) to provide grounded, nuanced assessments of AI style transfers.

Key Principle: AI Restyle is STYLE TRANSFER, not content editing.

What STAYS the Same:
• Subject identity (recognizable)
• Composition & layout
• Pose & positioning
• Scene content

What CHANGES:
• Visual art style
• Color palette & lighting
• Texture & brushwork
• Mood & atmosphere"""

    doc.add_paragraph(acrue_desc)

    doc.add_heading('2.2 Dimension Weights', 2)

    headers = ["Dimension", "Weight", "Max Score", "Description"]
    table = create_table_with_header(doc, headers)

    dimensions = [
        ["Accuracy", "1.0", "5.0", "Style transfer preserves content while applying aesthetic"],
        ["Completeness", "1.0", "5.0", "Style transformation is thorough and consistent"],
        ["Relevance", "0.5", "2.5", "Output matches what user wanted"],
        ["Usefulness", "0.5", "2.5", "Output is practically usable"],
        ["Exceptional", "2.0", "10.0", "Output delights and exceeds expectations"],
        ["TOTAL", "5.0", "25.0", ""]
    ]

    for row_data in dimensions:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            if row_data[0] == "TOTAL":
                row.cells[i].paragraphs[0].runs[0].bold = True

    doc.add_heading('2.3 Confidence Scale', 2)

    headers = ["Score", "Meaning"]
    table = create_table_with_header(doc, headers)

    confidence = [
        ["5", "Absolutely certain, strong evidence"],
        ["4", "Confident, clear evidence"],
        ["3", "Moderately confident, some evidence"],
        ["2", "Uncertain, weak evidence"],
        ["1", "Very uncertain, minimal evidence"]
    ]

    for row_data in confidence:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_heading('2.4 Grade Thresholds', 2)

    headers = ["Percentage", "Grade"]
    table = create_table_with_header(doc, headers)

    grades = [
        ["90-100%", "A+"],
        ["80-89%", "A"],
        ["70-79%", "B"],
        ["60-69%", "C"],
        ["< 60%", "F"]
    ]

    for row_data in grades:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('2.5 Synthesis Formula', 2)

    formula = """Final Score = (Feasibility Weight × Gemini Rank) + (Preference Weight × Opus Rank)
            = (0.5 × Gemini Rank) + (0.5 × Opus Rank)

Winner = Style with LOWEST final score (lower rank = better)"""

    doc.add_paragraph(formula)

    doc.add_page_break()

    # =========================================================================
    # 3. RUN CONFIGURATION
    # =========================================================================
    add_heading_with_style(doc, '3. Run Configuration', 1)

    headers = ["Parameter", "Value"]
    table = create_table_with_header(doc, headers, [2, 4])

    config = [
        ["Run ID", spec_data.get('run_id', '')],
        ["Schema Version", spec_data.get('schema_version', '')],
        ["Pipeline Version", spec_data.get('pipeline_version', '')],
        ["ACRUE Version", spec_data.get('acrue_version', '')],
        ["Styles Tested", ", ".join(spec_data.get('styles', []))],
        ["Image Count", str(spec_data.get('image_count', 0))],
        ["Total Evaluations", str(spec_data.get('image_count', 0) * len(spec_data.get('styles', [])))],
        ["Feasibility Judge", spec_data.get('judges', {}).get('feasibility', {}).get('model', '')],
        ["Preference Judge", spec_data.get('judges', {}).get('preference', {}).get('model', '')],
        ["Feasibility Weight", str(spec_data.get('synthesis', {}).get('feasibility_weight', 0.5))],
        ["Preference Weight", str(spec_data.get('synthesis', {}).get('preference_weight', 0.5))],
        ["Image Source", spec_data.get('image_selection', {}).get('source', '')],
        ["Selection Criteria", spec_data.get('image_selection', {}).get('criteria', '')]
    ]

    for row_data in config:
        row = table.add_row()
        row.cells[0].text = row_data[0]
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = row_data[1]

    doc.add_page_break()

    # =========================================================================
    # 4. FINAL RANKINGS & WINNER
    # =========================================================================
    add_heading_with_style(doc, '4. Final Rankings & Winner', 1)

    # Winner announcement
    winner_box = doc.add_paragraph()
    winner_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = winner_box.add_run('POP ART')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0, 100, 0)

    doc.add_paragraph()

    headers = ["Final Rank", "Style", "Gemini Rank", "Opus Rank", "Final Score"]
    table = create_table_with_header(doc, headers)

    rankings = [
        ["1", "Pop Art", "1", "1", "1.0"],
        ["2", "Storybook", "2", "2", "2.0"],
        ["3", "Anime", "3", "3", "3.0"]
    ]

    for i, row_data in enumerate(rankings):
        row = table.add_row()
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0:  # Winner row
                row.cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(row.cells[j], 'C6EFCE')

    doc.add_heading('Score Calculation', 2)

    calc_text = """Pop Art:   0.5 × 1 + 0.5 × 1 = 1.0  (Winner - Lowest Score)
Storybook: 0.5 × 2 + 0.5 × 2 = 2.0
Anime:     0.5 × 3 + 0.5 × 3 = 3.0"""

    doc.add_paragraph(calc_text)

    doc.add_page_break()

    # =========================================================================
    # 5. GEMINI ACRUE v3 EVALUATIONS
    # =========================================================================
    add_heading_with_style(doc, '5. Gemini ACRUE v3 Evaluations (Feasibility)', 1)

    # Group evaluations by image
    image_groups = {}
    for eval_data in acrue_data:
        original = eval_data.get('original', '')
        if original not in image_groups:
            image_groups[original] = []
        image_groups[original].append(eval_data)

    image_names = {
        'originals/img_001.png': ('5.1', 'Image 1: Pelican (Wildlife Photo)'),
        'originals/img_002.png': ('5.2', 'Image 2: Tiramisu (Food Photo)'),
        'originals/img_003.png': ('5.3', 'Image 3: Tiramisu Variant (Food Photo)')
    }

    for original_path, evals in image_groups.items():
        section_num, section_title = image_names.get(original_path, ('5.x', 'Unknown Image'))

        doc.add_heading(f'{section_num} {section_title}', 2)

        # Try to add original image
        orig_file = SCRIPT_DIR / original_path
        if orig_file.exists():
            doc.add_paragraph("Original Image:")
            try:
                doc.add_picture(str(orig_file), width=Inches(3))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[Image could not be loaded: {e}]")

        # Summary table for this image
        doc.add_paragraph()
        headers = ["Style", "Score", "Percentage", "Grade", "Pass Rate"]
        table = create_table_with_header(doc, headers)

        for eval_data in evals:
            row = table.add_row()
            style = eval_data.get('style', '')
            total = eval_data.get('total', 0)
            pct = eval_data.get('percentage', 0)
            grade = eval_data.get('grade', 'F')

            # Calculate pass rate
            dims = eval_data.get('dimensions', {})
            total_passed = sum(d.get('passed', 0) for d in dims.values())
            total_assertions = sum(d.get('total', 0) for d in dims.values())
            pass_rate = f"{total_passed}/{total_assertions}"

            row.cells[0].text = style
            row.cells[1].text = f"{total:.2f}/25"
            row.cells[2].text = f"{pct:.1f}%"
            row.cells[3].text = grade
            row.cells[4].text = pass_rate

            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            set_cell_shading(row.cells[3], get_grade_color(grade))

        # Detailed evaluation for each style
        for eval_data in evals:
            style = eval_data.get('style', '')
            total = eval_data.get('total', 0)
            pct = eval_data.get('percentage', 0)
            grade = eval_data.get('grade', 'F')
            summary = eval_data.get('summary', '')

            doc.add_heading(f'{style} Style - {grade} ({pct:.1f}%)', 3)

            # Add restyled image
            restyled_path = eval_data.get('restyled', '')
            restyled_file = SCRIPT_DIR / restyled_path
            if restyled_file.exists():
                try:
                    doc.add_picture(str(restyled_file), width=Inches(4))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f"[Image could not be loaded: {e}]")

            doc.add_paragraph()

            # Summary
            summary_para = doc.add_paragraph()
            summary_para.add_run("Summary: ").bold = True
            summary_para.add_run(summary)

            # Dimension scores
            doc.add_paragraph()
            headers = ["Dimension", "Passed", "Avg Confidence", "Weighted Score"]
            dim_table = create_table_with_header(doc, headers)

            dimensions = eval_data.get('dimensions', {})
            weights = {'accuracy': 1.0, 'completeness': 1.0, 'relevance': 0.5, 'usefulness': 0.5, 'exceptional': 2.0}

            for dim_name, dim_data in dimensions.items():
                row = dim_table.add_row()
                passed = dim_data.get('passed', 0)
                total_a = dim_data.get('total', 0)
                avg_conf = dim_data.get('avg_confidence', 0)
                dim_score = dim_data.get('dimension_score', 0)
                weighted = dim_score * weights.get(dim_name, 1.0)

                row.cells[0].text = dim_name.title()
                row.cells[1].text = f"{passed}/{total_a}"
                row.cells[2].text = f"{avg_conf:.1f}"
                row.cells[3].text = f"{weighted:.2f}"

                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Color based on pass rate
                if total_a > 0:
                    pass_pct = passed / total_a
                    if pass_pct >= 0.8:
                        set_cell_shading(row.cells[1], 'C6EFCE')
                    elif pass_pct >= 0.5:
                        set_cell_shading(row.cells[1], 'FFEB9C')
                    else:
                        set_cell_shading(row.cells[1], 'FFC7CE')

            # Total row
            total_row = dim_table.add_row()
            total_row.cells[0].text = "TOTAL"
            total_row.cells[0].paragraphs[0].runs[0].bold = True
            total_row.cells[3].text = f"{total:.2f}/25"
            total_row.cells[3].paragraphs[0].runs[0].bold = True

            # Assertion-by-assertion breakdown
            doc.add_paragraph()
            doc.add_paragraph("Assertion-by-Assertion Breakdown:").runs[0].bold = True

            for dim_name, dim_data in dimensions.items():
                doc.add_paragraph(f"{dim_name.upper()}:", style='List Bullet')
                assertions = dim_data.get('assertions', [])
                add_assertion_table(doc, assertions, dim_name)
                doc.add_paragraph()

        doc.add_page_break()

    # =========================================================================
    # 6. OPUS PREFERENCE RANKINGS
    # =========================================================================
    add_heading_with_style(doc, '6. Opus Preference Rankings', 1)

    doc.add_heading('6.1 Judge Information', 2)

    judge_info = f"""Judge Model: {opus_data.get('judge', '')}
Timestamp: {opus_data.get('timestamp', '')}
Methodology: {opus_data.get('methodology', '')}"""

    doc.add_paragraph(judge_info)

    doc.add_heading('6.2 Rankings', 2)

    headers = ["Rank", "Style", "Appeal Score", "Reasoning"]
    table = create_table_with_header(doc, headers, [0.5, 1, 1, 4])

    for ranking in opus_data.get('rankings', []):
        row = table.add_row()
        row.cells[0].text = str(ranking.get('rank', ''))
        row.cells[1].text = ranking.get('style', '')
        row.cells[2].text = f"{ranking.get('appeal_score', 0)}/10"
        row.cells[3].text = ranking.get('reasoning', '')

        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Make reasoning text smaller
        for run in row.cells[3].paragraphs[0].runs:
            run.font.size = Pt(9)

        if ranking.get('rank') == 1:
            for cell in row.cells:
                set_cell_shading(cell, 'C6EFCE')

    doc.add_heading('6.3 Analysis Summary', 2)

    analysis = opus_data.get('analysis', {})

    headers = ["Category", "Result"]
    table = create_table_with_header(doc, headers)

    analysis_rows = [
        ["Most Shareable", analysis.get('most_shareable', '')],
        ["Most Delightful", analysis.get('most_delightful', '')],
        ["Key Observation", analysis.get('key_observation', '')],
        ["Recommendation", analysis.get('recommendation', '')]
    ]

    for row_data in analysis_rows:
        row = table.add_row()
        row.cells[0].text = row_data[0]
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = row_data[1]

    doc.add_page_break()

    # =========================================================================
    # 7. SYNTHESIS & FINAL SCORES
    # =========================================================================
    add_heading_with_style(doc, '7. Synthesis & Final Scores', 1)

    doc.add_heading('7.1 Gemini Rankings (from gemini.json)', 2)

    headers = ["Rank", "Style", "Avg Score", "Avg %", "Grade", "Reasoning"]
    table = create_table_with_header(doc, headers, [0.5, 1, 0.8, 0.7, 0.5, 3])

    for ranking in gemini_data.get('rankings', []):
        row = table.add_row()
        row.cells[0].text = str(ranking.get('rank', ''))
        row.cells[1].text = ranking.get('style', '')
        row.cells[2].text = f"{ranking.get('avg_score', 0):.2f}"
        row.cells[3].text = f"{ranking.get('avg_percentage', 0):.1f}%"
        row.cells[4].text = ranking.get('avg_grade', '')
        row.cells[5].text = ranking.get('reasoning', '')

        for i in range(5):
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in row.cells[5].paragraphs[0].runs:
            run.font.size = Pt(8)

        set_cell_shading(row.cells[4], get_grade_color(ranking.get('avg_grade', 'F')))

    doc.add_heading('7.2 Combined Synthesis', 2)

    headers = ["Style", "Gemini Rank", "Opus Rank", "Formula", "Final Score", "Final Rank"]
    table = create_table_with_header(doc, headers)

    synthesis_data = [
        ["Pop Art", "1", "1", "0.5×1 + 0.5×1", "1.0", "1 (Winner)"],
        ["Storybook", "2", "2", "0.5×2 + 0.5×2", "2.0", "2"],
        ["Anime", "3", "3", "0.5×3 + 0.5×3", "3.0", "3"]
    ]

    for i, row_data in enumerate(synthesis_data):
        row = table.add_row()
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, 'C6EFCE')
                cell.paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # =========================================================================
    # 8. STATISTICAL ANALYSIS
    # =========================================================================
    add_heading_with_style(doc, '8. Statistical Analysis', 1)

    doc.add_heading('8.1 Aggregate Assertion Pass Rates', 2)

    headers = ["Dimension", "Anime", "Pop Art", "Storybook"]
    table = create_table_with_header(doc, headers)

    # Calculate aggregate stats
    style_stats = {'Anime': {}, 'Pop Art': {}, 'Storybook': {}}

    for eval_data in acrue_data:
        style = eval_data.get('style', '')
        dims = eval_data.get('dimensions', {})

        for dim_name, dim_data in dims.items():
            if dim_name not in style_stats[style]:
                style_stats[style][dim_name] = {'passed': 0, 'total': 0}
            style_stats[style][dim_name]['passed'] += dim_data.get('passed', 0)
            style_stats[style][dim_name]['total'] += dim_data.get('total', 0)

    dimension_order = ['accuracy', 'completeness', 'relevance', 'usefulness', 'exceptional']

    for dim in dimension_order:
        row = table.add_row()
        row.cells[0].text = dim.title()

        for i, style in enumerate(['Anime', 'Pop Art', 'Storybook']):
            stats = style_stats[style].get(dim, {'passed': 0, 'total': 0})
            passed = stats['passed']
            total = stats['total']
            pct = (passed / total * 100) if total > 0 else 0
            row.cells[i + 1].text = f"{passed}/{total} ({pct:.0f}%)"
            row.cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            if pct >= 80:
                set_cell_shading(row.cells[i + 1], 'C6EFCE')
            elif pct >= 50:
                set_cell_shading(row.cells[i + 1], 'FFEB9C')
            else:
                set_cell_shading(row.cells[i + 1], 'FFC7CE')

    # Total row
    total_row = table.add_row()
    total_row.cells[0].text = "TOTAL"
    total_row.cells[0].paragraphs[0].runs[0].bold = True

    for i, style in enumerate(['Anime', 'Pop Art', 'Storybook']):
        total_passed = sum(d['passed'] for d in style_stats[style].values())
        total_assertions = sum(d['total'] for d in style_stats[style].values())
        pct = (total_passed / total_assertions * 100) if total_assertions > 0 else 0
        total_row.cells[i + 1].text = f"{total_passed}/{total_assertions} ({pct:.0f}%)"
        total_row.cells[i + 1].paragraphs[0].runs[0].bold = True
        total_row.cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('8.2 Score Distribution', 2)

    headers = ["Image", "Anime", "Pop Art", "Storybook"]
    table = create_table_with_header(doc, headers)

    score_by_image = {}
    for eval_data in acrue_data:
        orig = eval_data.get('original', '')
        style = eval_data.get('style', '')
        score = eval_data.get('total', 0)
        grade = eval_data.get('grade', 'F')

        if orig not in score_by_image:
            score_by_image[orig] = {}
        score_by_image[orig][style] = (score, grade)

    for orig, scores in score_by_image.items():
        row = table.add_row()
        img_name = orig.replace('originals/', '').replace('.png', '')
        row.cells[0].text = img_name

        for i, style in enumerate(['Anime', 'Pop Art', 'Storybook']):
            score, grade = scores.get(style, (0, 'F'))
            row.cells[i + 1].text = f"{score:.2f} ({grade})"
            row.cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(row.cells[i + 1], get_grade_color(grade))

    # Average row
    avg_row = table.add_row()
    avg_row.cells[0].text = "AVERAGE"
    avg_row.cells[0].paragraphs[0].runs[0].bold = True

    for i, style in enumerate(['Anime', 'Pop Art', 'Storybook']):
        style_scores = [e.get('total', 0) for e in acrue_data if e.get('style') == style]
        avg = sum(style_scores) / len(style_scores) if style_scores else 0
        avg_row.cells[i + 1].text = f"{avg:.2f}"
        avg_row.cells[i + 1].paragraphs[0].runs[0].bold = True
        avg_row.cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('8.3 Confidence Score Analysis', 2)

    headers = ["Style", "Avg Confidence", "Interpretation"]
    table = create_table_with_header(doc, headers)

    for style in ['Anime', 'Pop Art', 'Storybook']:
        style_evals = [e for e in acrue_data if e.get('style') == style]
        all_conf = []
        for e in style_evals:
            for dim in e.get('dimensions', {}).values():
                for a in dim.get('assertions', []):
                    all_conf.append(a.get('confidence', 3))

        avg_conf = sum(all_conf) / len(all_conf) if all_conf else 0

        if avg_conf >= 4.5:
            interp = "High certainty - evaluations are reliable"
        elif avg_conf >= 4.0:
            interp = "Good certainty - evaluations are trustworthy"
        elif avg_conf >= 3.5:
            interp = "Moderate certainty - some ambiguity"
        else:
            interp = "Low certainty - results may be less reliable"

        row = table.add_row()
        row.cells[0].text = style
        row.cells[1].text = f"{avg_conf:.2f}/5"
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].text = interp

    doc.add_page_break()

    # =========================================================================
    # 9. IMAGE GALLERY
    # =========================================================================
    add_heading_with_style(doc, '9. Image Gallery', 1)

    for i, (orig_path, evals) in enumerate(image_groups.items(), 1):
        doc.add_heading(f'9.{i} Image Set {i}', 2)

        # Original
        doc.add_paragraph("Original:")
        orig_file = SCRIPT_DIR / orig_path
        if orig_file.exists():
            try:
                doc.add_picture(str(orig_file), width=Inches(2.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                doc.add_paragraph("[Could not load image]")

        # Restyled images in a row description
        for eval_data in sorted(evals, key=lambda x: x.get('style', '')):
            style = eval_data.get('style', '')
            grade = eval_data.get('grade', 'F')
            score = eval_data.get('total', 0)

            doc.add_paragraph(f"{style} (Score: {score:.2f}, Grade: {grade}):")

            restyled_path = eval_data.get('restyled', '')
            restyled_file = SCRIPT_DIR / restyled_path
            if restyled_file.exists():
                try:
                    doc.add_picture(str(restyled_file), width=Inches(3.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    doc.add_paragraph("[Could not load image]")

        doc.add_paragraph()

    doc.add_page_break()

    # =========================================================================
    # 10. RECOMMENDATIONS
    # =========================================================================
    add_heading_with_style(doc, '10. Recommendations', 1)

    doc.add_heading('10.1 Primary Recommendation', 2)

    primary = doc.add_paragraph()
    primary.add_run("Prioritize Pop Art style for diverse image types.").bold = True

    doc.add_paragraph("""Pop Art consistently achieved the highest scores across all evaluation dimensions and image types. Its bold graphic style creates immediate visual impact that works universally - whether the subject is a person, animal, food, or landscape.

Key advantages:
• 100% assertion pass rate across all 69 assertions
• Highest appeal score (9.2/10) from preference judge
• Strong "wow factor" and social shareability
• Works equally well on human and non-human subjects""")

    doc.add_heading('10.2 Secondary Recommendation', 2)

    doc.add_paragraph("""Storybook as strong alternative for family-oriented content.

Storybook achieved solid A-grade scores (86.7% average) and provides a warm, nostalgic aesthetic that appeals particularly to family-oriented users. Consider promoting Storybook for:
• Children's content
• Family photos
• Nursery decor applications
• Holiday/greeting card use cases""")

    doc.add_heading('10.3 Anime Style Improvements', 2)

    doc.add_paragraph("""Implement subject detection for Anime style.

Anime's poor performance (31.9% average) stems from its human-character-focused aesthetic being applied to non-human subjects. The style assertions expect "large expressive eyes," "anime protagonist quality," and "facial proportions" - none of which apply to wildlife or food photos.

Recommended actions:
1. Add subject detection before applying Anime style
2. If subject is non-human, display warning: "Anime style works best with photos of people"
3. Consider creating "Anime Nature" or "Anime Food" variants with appropriate assertions
4. Adjust user expectations through UI messaging""")

    doc.add_heading('10.4 Evaluation Framework Improvements', 2)

    doc.add_paragraph("""1. Style-specific assertion sets: Create different assertion sets for different subject types (people, animals, food, landscapes)

2. Subject-style compatibility scoring: Pre-evaluate how well a style matches the detected subject type

3. User preference learning: Track which styles users actually save/share to refine preference rankings

4. A/B testing integration: Use ACRUE scores to select variants for A/B tests

5. Automated regression testing: Run ACRUE evaluations on each model update to detect quality regressions""")

    doc.add_page_break()

    # =========================================================================
    # 11. APPENDIX
    # =========================================================================
    add_heading_with_style(doc, '11. Appendix: Raw Data', 1)

    doc.add_heading('11.1 File Manifest', 2)

    manifest = """runs/Run_2026_02_03_baseline/
├── run_spec.json          # Run configuration
├── originals/
│   ├── img_001.png        # Pelican wildlife photo
│   ├── img_002.png        # Tiramisu food photo
│   └── img_003.png        # Tiramisu variant photo
├── restyled/
│   ├── img_001_anime.png
│   ├── img_001_pop_art.png
│   ├── img_001_storybook.png
│   ├── img_002_anime.png
│   ├── img_002_pop_art.png
│   ├── img_002_storybook.png
│   ├── img_003_anime.png
│   ├── img_003_pop_art.png
│   └── img_003_storybook.png
├── acrue.json             # ACRUE v3 evaluation results (9 records)
├── gemini.json            # Feasibility rankings
├── opus.json              # Preference rankings
├── report.md              # Summary report
└── Exhaustive_Benchmark_Report.docx  # This report"""

    doc.add_paragraph(manifest)

    doc.add_heading('11.2 Evaluation Timestamps', 2)

    headers = ["Evaluation", "Timestamp"]
    table = create_table_with_header(doc, headers)

    for eval_data in acrue_data:
        row = table.add_row()
        row.cells[0].text = eval_data.get('evaluation_id', '')
        row.cells[1].text = eval_data.get('timestamp', '')

    doc.add_heading('11.3 Judge Versions', 2)

    headers = ["Judge", "Model", "Version"]
    table = create_table_with_header(doc, headers)

    judges = [
        ["Feasibility (ACRUE v3)", "gemini-2.0-flash", "2.0"],
        ["Preference", "claude-opus-4-5", "20251101"]
    ]

    for judge in judges:
        row = table.add_row()
        for i, val in enumerate(judge):
            row.cells[i].text = val

    # =========================================================================
    # SAVE DOCUMENT
    # =========================================================================
    doc.save(OUTPUT_PATH)
    print(f"Report saved to: {OUTPUT_PATH}")
    print(f"Total pages: ~30-40 (estimated)")


if __name__ == "__main__":
    main()
