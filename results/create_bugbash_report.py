from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pathlib import Path
import os

# Screenshot folder
screenshot_folder = str(Path(__file__).resolve().parent.parent.parent / '.playwright-mcp')

# Create document
doc = Document()

# Title
title = doc.add_heading('AI Restyle Feature - Bug Bash Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run('Test Date: ').bold = True
meta.add_run('January 27, 2026\n')
meta.add_run('Tester: ').bold = True
meta.add_run('Automated QA Agent (Claude Code with Playwright MCP)\n')
meta.add_run('Feature: ').bold = True
meta.add_run('OneDrive Photos - AI Restyle\n')
meta.add_run('Test Environment: ').bold = True
meta.add_run('OneDrive Web (onedrive.live.com)')

# Executive Summary
doc.add_heading('Executive Summary', level=1)

summary_table = doc.add_table(rows=5, cols=2)
summary_table.style = 'Table Grid'
summary_data = [
    ('Metric', 'Value'),
    ('Total Tests', '12'),
    ('Passed', '10'),
    ('Bugs Found', '2'),
    ('Severity', '1 High, 1 Medium')
]
for i, (col1, col2) in enumerate(summary_data):
    summary_table.rows[i].cells[0].text = col1
    summary_table.rows[i].cells[1].text = col2
    if i == 0:
        summary_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        summary_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

def add_screenshot(doc, filename, caption):
    """Add screenshot to document if it exists"""
    filepath = os.path.join(screenshot_folder, filename)
    if os.path.exists(filepath):
        doc.add_paragraph()
        doc.add_picture(filepath, width=Inches(6))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        return True
    else:
        doc.add_paragraph(f'[Screenshot not found: {filename}]')
        return False

# ============== BUGS FOUND ==============
doc.add_heading('Bugs Found', level=1)

# Bug #1
doc.add_heading('BUG #1: WEBP Format Not Supported for AI Restyle', level=2)

# Severity badge
p = doc.add_paragraph()
p.add_run('Severity: ').bold = True
run = p.add_run('MEDIUM')
run.font.color.rgb = RGBColor(255, 165, 0)
run.bold = True
p.add_run('  |  ')
p.add_run('Priority: ').bold = True
p.add_run('P2')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Description:\n').bold = True
p.add_run('WEBP format images do not have the "Edit" or "Restyle with AI" buttons available in the image viewer toolbar. Only basic options are shown.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Steps to Reproduce:\n').bold = True
doc.add_paragraph('1. Navigate to OneDrive Photos gallery', style='List Number')
doc.add_paragraph('2. Open a .webp format image (e.g., OIP.webp)', style='List Number')
doc.add_paragraph('3. Observe the toolbar options', style='List Number')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Expected Behavior:\n').bold = True
p.add_run('WEBP images should have the same editing options as PNG/JPEG, including "Restyle with AI"')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Actual Behavior:\n').bold = True
p.add_run('WEBP images only show: Close, Delete, Download, Add to album, Share, Favorite\n')
p.add_run('Missing: ').bold = True
p.add_run('Edit, Restyle with AI, Edit with Designer')

add_screenshot(doc, 'bugbash_01_webp_no_restyle.png', 'Figure 1: WEBP image toolbar - Missing Restyle with AI button')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Impact:\n').bold = True
p.add_run('Users cannot use AI Restyle on WEBP images, which is a common modern image format. This may cause confusion as users won\'t understand why the feature is unavailable.')

# Bug #2
doc.add_heading('BUG #2: Stop Button Does Not Cancel Generation', level=2)

# Severity badge
p = doc.add_paragraph()
p.add_run('Severity: ').bold = True
run = p.add_run('HIGH')
run.font.color.rgb = RGBColor(255, 0, 0)
run.bold = True
p.add_run('  |  ')
p.add_run('Priority: ').bold = True
p.add_run('P1')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Description:\n').bold = True
p.add_run('The Stop button displayed during AI generation does not actually cancel the generation process. Multiple clicks on the Stop button are ignored, and the generation continues to completion.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Steps to Reproduce:\n').bold = True
doc.add_paragraph('1. Open a supported image (PNG/JPEG)', style='List Number')
doc.add_paragraph('2. Click "Restyle with AI"', style='List Number')
doc.add_paragraph('3. Select any style (e.g., Pop Art)', style='List Number')
doc.add_paragraph('4. Click Send to start generation', style='List Number')
doc.add_paragraph('5. While loading message shows, click Stop button', style='List Number')
doc.add_paragraph('6. Click Stop button multiple times (tested 3 clicks)', style='List Number')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Expected Behavior:\n').bold = True
p.add_run('Generation should stop immediately or within a few seconds, returning to the original image state.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Actual Behavior:\n').bold = True
doc.add_paragraph('Stop button clicked 3 times during generation', style='List Bullet')
doc.add_paragraph('Loading messages continued cycling through different states', style='List Bullet')
doc.add_paragraph('Generation completed successfully despite Stop attempts', style='List Bullet')
doc.add_paragraph('Generated image was produced as if Stop was never clicked', style='List Bullet')

add_screenshot(doc, 'bugbash_02_stop_not_working.png', 'Figure 2: Generation completed despite clicking Stop button 3 times')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Impact:\n').bold = True
doc.add_paragraph('Users cannot cancel accidental generations', style='List Bullet')
doc.add_paragraph('Users waste time waiting for unwanted generations to complete', style='List Bullet')
doc.add_paragraph('Poor UX as the Stop button implies functionality that doesn\'t work', style='List Bullet')

# ============== TEST SCENARIOS ==============
doc.add_heading('Test Scenarios', level=1)

# Scenario 1
doc.add_heading('Scenario 1: Format Support Testing', level=2)
p = doc.add_paragraph()
p.add_run('Objective: ').bold = True
p.add_run('Verify which image formats support AI Restyle')

results_table = doc.add_table(rows=4, cols=3)
results_table.style = 'Table Grid'
format_data = [
    ('Format', 'Restyle Available', 'Status'),
    ('PNG', 'Yes - Full toolbar', 'PASS'),
    ('JPEG/JPG', 'Yes - Full toolbar', 'PASS'),
    ('WEBP', 'No - Missing buttons', 'BUG')
]
for i, row_data in enumerate(format_data):
    for j, cell_text in enumerate(row_data):
        results_table.rows[i].cells[j].text = cell_text
        if i == 0:
            results_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

# Scenario 2
doc.add_heading('Scenario 2: Rapid Style Switching', level=2)
p = doc.add_paragraph()
p.add_run('Status: ').bold = True
run = p.add_run('PASS')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True

p = doc.add_paragraph()
p.add_run('Test: ').bold = True
p.add_run('Quickly switch between multiple styles (Movie Poster -> Anime -> Pop Art)')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Result: ').bold = True
p.add_run('Prompt text updates correctly each time without errors or lag. No race conditions observed.')

# Scenario 3
doc.add_heading('Scenario 3: Double-Click on Send Button', level=2)
p = doc.add_paragraph()
p.add_run('Status: ').bold = True
run = p.add_run('PASS')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True

p = doc.add_paragraph()
p.add_run('Test: ').bold = True
p.add_run('Double-click the Send button to check for duplicate generation requests')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Result: ').bold = True
p.add_run('Only one generation is triggered. UI handles double-click gracefully without creating duplicate requests.')

# Scenario 4
doc.add_heading('Scenario 4: Stop Button During Generation', level=2)
p = doc.add_paragraph()
p.add_run('Status: ').bold = True
run = p.add_run('FAIL - BUG #2')
run.font.color.rgb = RGBColor(255, 0, 0)
run.bold = True

p = doc.add_paragraph()
p.add_run('Test: ').bold = True
p.add_run('Click Stop button during AI generation')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Result: ').bold = True
p.add_run('Stop button does not cancel generation. See Bug #2 for details.')

# Scenario 5
doc.add_heading('Scenario 5: Unsaved Changes Dialog', level=2)
p = doc.add_paragraph()
p.add_run('Status: ').bold = True
run = p.add_run('PASS')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True

p = doc.add_paragraph()
p.add_run('Test: ').bold = True
p.add_run('Click Back button after generating an image without saving')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Result: ').bold = True
p.add_run('Confirmation dialog appears with clear messaging:')
doc.add_paragraph('"Leave without saving?"', style='List Bullet')
doc.add_paragraph('"You haven\'t saved your photo yet. Going back now will remove your progress so far."', style='List Bullet')
doc.add_paragraph('Options: Save, Discard', style='List Bullet')

add_screenshot(doc, 'bugbash_03_unsaved_dialog.png', 'Figure 3: Unsaved changes confirmation dialog - Good UX')

# Scenario 6
doc.add_heading('Scenario 6: Discard Functionality', level=2)
p = doc.add_paragraph()
p.add_run('Status: ').bold = True
run = p.add_run('PASS')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True

p = doc.add_paragraph()
p.add_run('Test: ').bold = True
p.add_run('Click Discard in the unsaved changes dialog')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Result: ').bold = True
p.add_run('Generated image is discarded, original image is restored, and editor closes correctly.')

# Scenario 7
doc.add_heading('Scenario 7: UI Button State Management', level=2)
p = doc.add_paragraph()
p.add_run('Status: ').bold = True
run = p.add_run('PASS')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True

p = doc.add_paragraph()
p.add_run('Test: ').bold = True
p.add_run('Verify button states after generation completes')

state_table = doc.add_table(rows=6, cols=2)
state_table.style = 'Table Grid'
state_data = [
    ('Button', 'State After Generation'),
    ('Undo', 'Enabled'),
    ('Redo', 'Disabled'),
    ('Reset', 'Enabled'),
    ('Download', 'Enabled'),
    ('Save copy', 'Enabled')
]
for i, row_data in enumerate(state_data):
    for j, cell_text in enumerate(row_data):
        state_table.rows[i].cells[j].text = cell_text
        if i == 0:
            state_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

# Scenario 8
doc.add_heading('Scenario 8: Custom Prompt Entry', level=2)
p = doc.add_paragraph()
p.add_run('Status: ').bold = True
run = p.add_run('PASS')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True

p = doc.add_paragraph()
p.add_run('Test: ').bold = True
p.add_run('Clear preset prompt and enter custom text')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Result: ').bold = True
p.add_run('Successfully entered custom prompt: "Transform into a watercolor painting with soft pastel colors". Send button remained enabled.')

# Scenario 9
doc.add_heading('Scenario 9: All 14 Style Presets Verification', level=2)
p = doc.add_paragraph()
p.add_run('Status: ').bold = True
run = p.add_run('PASS')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True

p = doc.add_paragraph()
p.add_run('All 14 presets verified present:\n').bold = True

presets_table = doc.add_table(rows=5, cols=3)
presets_table.style = 'Table Grid'
presets_data = [
    ('Movie Poster', 'Plush Toy', 'Anime'),
    ('Chibi Sticker', 'Caricature', 'Superhero'),
    ('Toy Model', 'Graffiti', 'Crochet Art'),
    ('Doodle', 'Pencil Portrait', 'Storybook'),
    ('Photo Booth', 'Pop Art', '')
]
for i, row_data in enumerate(presets_data):
    for j, cell_text in enumerate(row_data):
        presets_table.rows[i].cells[j].text = cell_text

# ============== RECOMMENDATIONS ==============
doc.add_heading('Recommendations', level=1)

doc.add_heading('High Priority', level=2)
doc.add_paragraph('Fix Stop Button - Implement actual cancellation of AI generation when Stop is clicked, or remove the button if cancellation is not technically feasible', style='List Bullet')

doc.add_heading('Medium Priority', level=2)
doc.add_paragraph('Add WEBP Support - Enable AI Restyle for WEBP format images', style='List Bullet')
doc.add_paragraph('Format Messaging - If WEBP support cannot be added, show a tooltip or message explaining which formats are supported', style='List Bullet')

doc.add_heading('Low Priority', level=2)
doc.add_paragraph('Add progress indicator with estimated time remaining during generation', style='List Bullet')
doc.add_paragraph('Add keyboard shortcuts (Ctrl+Z/Ctrl+Y) for Undo/Redo', style='List Bullet')
doc.add_paragraph('Add explicit Copy to Clipboard button for generated images', style='List Bullet')

# ============== CONCLUSION ==============
doc.add_heading('Conclusion', level=1)
conclusion = doc.add_paragraph()
conclusion.add_run('The AI Restyle feature is largely functional with a good user experience for supported image formats. ').bold = True
conclusion.add_run('The feature handles edge cases well (double-clicks, rapid switching, unsaved changes). However, two bugs were identified:\n\n')
conclusion.add_run('1. Stop button non-functional (High severity) - Users cannot cancel generations\n')
conclusion.add_run('2. WEBP format unsupported (Medium severity) - Modern format excluded\n\n')
conclusion.add_run('The Stop button issue should be addressed before production release as it affects user control over the feature and creates a misleading UI element.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Report Generated: ').bold = True
p.add_run('January 27, 2026\n')
p.add_run('Automation Tool: ').bold = True
p.add_run('Claude Code with Playwright MCP\n')
p.add_run('Total Test Duration: ').bold = True
p.add_run('~10 minutes')

# Save document
output_path = str(Path(__file__).resolve().parent / 'AI_Restyle_Bug_Bash_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
