from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import os

# Screenshot folder
screenshot_folder = str(Path(__file__).resolve().parent.parent.parent / '.playwright-mcp')

# Create document
doc = Document()

# Title
title = doc.add_heading('Comprehensive AI Restyle Bug Bash Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
doc.add_paragraph()
meta = doc.add_paragraph()
meta.add_run('Test Date: ').bold = True
meta.add_run('January 27, 2026\n')
meta.add_run('Tester: ').bold = True
meta.add_run('Claude Code Automated QA Agent\n')
meta.add_run('Feature: ').bold = True
meta.add_run('OneDrive Photos - AI Restyle\n')
meta.add_run('Test Focus: ').bold = True
meta.add_run('Latency, Accessibility, UI/UX, Design Improvements')

# Executive Summary
doc.add_heading('Executive Summary', level=1)

summary_table = doc.add_table(rows=6, cols=2)
summary_table.style = 'Table Grid'
summary_data = [
    ('Category', 'Findings'),
    ('Latency Issues', '0 Critical, 1 Observation'),
    ('Accessibility Bugs', '5 Issues Found'),
    ('UI/UX Issues', '3 Improvements Suggested'),
    ('Design Suggestions', '6 Recommendations'),
    ('Overall Rating', 'Good with Room for Improvement')
]
for i, (col1, col2) in enumerate(summary_data):
    summary_table.rows[i].cells[0].text = col1
    summary_table.rows[i].cells[1].text = col2
    if i == 0:
        summary_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        summary_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

def add_screenshot(doc, filename, caption):
    """Add screenshot to document if it exists"""
    filepath = os.path.join(screenshot_folder, filename)
    if os.path.exists(filepath):
        doc.add_paragraph()
        doc.add_picture(filepath, width=Inches(6))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        return True
    else:
        doc.add_paragraph(f'[Screenshot not found: {filename}]')
        return False

# ============== LATENCY TRACKING ==============
doc.add_heading('Section 1: Latency Tracking', level=1)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Methodology: ').bold = True
p.add_run('JavaScript timestamps captured at action start/end using Date.now()')

latency_table = doc.add_table(rows=5, cols=3)
latency_table.style = 'Table Grid'
latency_data = [
    ('Action', 'Latency', 'Assessment'),
    ('Gallery Page Load', '~911ms', 'GOOD - Under 1 second'),
    ('Restyle Panel Open', '~500ms', 'GOOD - Quick response'),
    ('AI Generation (Pop Art)', '50.5 seconds', 'ACCEPTABLE - Complex AI task'),
    ('Undo/Redo Operations', '<100ms', 'EXCELLENT - Instant')
]
for i, row_data in enumerate(latency_data):
    for j, cell_text in enumerate(row_data):
        latency_table.rows[i].cells[j].text = cell_text
        if i == 0:
            latency_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Observation: ').bold = True
p.add_run('AI generation time of ~50 seconds is expected for complex image transformations. However, there is no progress indicator or estimated time remaining, which could improve user experience during the wait.')

add_screenshot(doc, 'audit_05_generation_started.png', 'Figure 1: Generation in progress - Loading messages cycle through different states')

# ============== ACCESSIBILITY AUDIT ==============
doc.add_heading('Section 2: Accessibility Audit', level=1)

# Bug A1
doc.add_heading('A11Y-BUG #1: Gallery Images Missing Alt Text', level=2)
p = doc.add_paragraph()
p.add_run('Severity: ').bold = True
run = p.add_run('HIGH')
run.font.color.rgb = RGBColor(255, 0, 0)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run('48 images in the Photos gallery lack descriptive alt text. Screen reader users cannot identify image content.')

p = doc.add_paragraph()
p.add_run('WCAG Violation: ').bold = True
p.add_run('1.1.1 Non-text Content (Level A)')

p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run('Implement AI-generated alt text descriptions for all gallery images, or use existing metadata (filename, tags, date) to generate meaningful descriptions.')

# Bug A2
doc.add_heading('A11Y-BUG #2: Missing Main Landmark in Gallery', level=2)
p = doc.add_paragraph()
p.add_run('Severity: ').bold = True
run = p.add_run('MEDIUM')
run.font.color.rgb = RGBColor(255, 165, 0)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run('The gallery page lacks a <main> landmark role on the primary content area (found later in image viewer).')

p = doc.add_paragraph()
p.add_run('WCAG Violation: ').bold = True
p.add_run('1.3.1 Info and Relationships (Level A)')

# Bug A3
doc.add_heading('A11Y-BUG #3: Style Presets Not Keyboard Navigable', level=2)
p = doc.add_paragraph()
p.add_run('Severity: ').bold = True
run = p.add_run('HIGH')
run.font.color.rgb = RGBColor(255, 0, 0)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run('Individual style presets in the Restyle panel cannot be navigated using arrow keys. Tab moves to the container, but arrow keys do not move between presets.')

p = doc.add_paragraph()
p.add_run('Steps to Reproduce:\n').bold = True
doc.add_paragraph('1. Open Restyle panel', style='List Number')
doc.add_paragraph('2. Press Tab to reach style presets container', style='List Number')
doc.add_paragraph('3. Press Arrow keys - focus does not move between individual presets', style='List Number')

p = doc.add_paragraph()
p.add_run('WCAG Violation: ').bold = True
p.add_run('2.1.1 Keyboard (Level A)')

p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run('Implement arrow key navigation between style presets using roving tabindex or aria-activedescendant pattern.')

add_screenshot(doc, 'audit_04_keyboard_nav_issue.png', 'Figure 2: Style presets container receives focus but individual items are not keyboard accessible')

# Bug A4
doc.add_heading('A11Y-BUG #4: Restyle Dialog Missing aria-label', level=2)
p = doc.add_paragraph()
p.add_run('Severity: ').bold = True
run = p.add_run('LOW')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run('The Restyle panel uses role="alertdialog" but lacks aria-label or aria-labelledby to announce its purpose.')

# Bug A5
doc.add_heading('A11Y-BUG #5: Some Buttons Missing aria-label', level=2)
p = doc.add_paragraph()
p.add_run('Severity: ').bold = True
run = p.add_run('LOW')
run.font.color.rgb = RGBColor(0, 128, 0)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run('"Save copy" and "Reset" buttons lack aria-label attributes. While they have visible text, explicit labels improve screen reader experience.')

# Accessibility Passes
doc.add_heading('Accessibility Passes', level=2)

pass_table = doc.add_table(rows=8, cols=2)
pass_table.style = 'Table Grid'
pass_data = [
    ('Check', 'Result'),
    ('Skip to main content link', 'PASS - Present and functional'),
    ('Toolbar buttons have aria-labels', 'PASS - All have descriptive labels'),
    ('Style preset images have alt text', 'PASS - All 14 presets labeled'),
    ('Unsaved changes dialog accessibility', 'PASS - Proper role, heading, focus trap'),
    ('Keyboard focus visible', 'PASS - Clear focus indicators'),
    ('Prompt textbox has aria-label', 'PASS - "Enter prompt"'),
    ('Color contrast (dark mode)', 'PASS - Sufficient contrast')
]
for i, row_data in enumerate(pass_data):
    for j, cell_text in enumerate(row_data):
        pass_table.rows[i].cells[j].text = cell_text
        if i == 0:
            pass_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

add_screenshot(doc, 'audit_01_keyboard_focus.png', 'Figure 3: Keyboard focus indicator visible on gallery images')

# ============== UI/UX CRITIQUE ==============
doc.add_heading('Section 3: UI/UX Critique', level=1)

# Issue 1
doc.add_heading('UX-ISSUE #1: No Generation Progress Indicator', level=2)
p = doc.add_paragraph()
p.add_run('Observation: ').bold = True
p.add_run('During the ~50 second AI generation, users see rotating messages ("Pixels getting warmed up...", "Getting your photo ready...") but no progress percentage or estimated time.')

p = doc.add_paragraph()
p.add_run('Impact: ').bold = True
p.add_run('Users cannot estimate how long to wait, leading to potential abandonment or frustration.')

p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run('Add a progress bar or percentage indicator. Even if exact time is unknown, showing progress stages (e.g., "Step 2 of 4") improves perceived performance.')

# Issue 2
doc.add_heading('UX-ISSUE #2: Stop Button Non-Functional (Known Bug)', level=2)
p = doc.add_paragraph()
p.add_run('Observation: ').bold = True
p.add_run('The Stop button appears during generation but does not cancel the operation (previously documented bug).')

p = doc.add_paragraph()
p.add_run('Impact: ').bold = True
p.add_run('Users cannot cancel accidental or unwanted generations, wasting time and resources.')

# Issue 3
doc.add_heading('UX-ISSUE #3: WEBP Format Unsupported (Known Bug)', level=2)
p = doc.add_paragraph()
p.add_run('Observation: ').bold = True
p.add_run('WEBP images do not show the Restyle option, with no explanation to users.')

p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run('Either add WEBP support or show a tooltip/message explaining format requirements.')

# Positive UX Elements
doc.add_heading('Positive UX Elements', level=2)

doc.add_paragraph('Unsaved changes dialog - Clear messaging and appropriate options (Save/Discard)', style='List Bullet')
doc.add_paragraph('Undo/Redo functionality - Works correctly with proper button state management', style='List Bullet')
doc.add_paragraph('Style presets - Visual previews help users understand expected output', style='List Bullet')
doc.add_paragraph('AI disclaimer - "AI might alter features, including faces" sets appropriate expectations', style='List Bullet')
doc.add_paragraph('Custom prompt support - Users can modify or write custom prompts', style='List Bullet')

add_screenshot(doc, 'audit_09_unsaved_dialog.png', 'Figure 4: Well-designed unsaved changes confirmation dialog')

# ============== DESIGN IMPROVEMENTS ==============
doc.add_heading('Section 4: Design Improvement Suggestions', level=1)

# Improvement 1
doc.add_heading('DESIGN-01: Add Progress Indicator During Generation', level=2)
p = doc.add_paragraph()
p.add_run('Current State: ').bold = True
p.add_run('Rotating text messages only')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Suggested Improvement: ').bold = True
p.add_run('Add a visual progress bar with percentage or step indicator')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Benefit: ').bold = True
p.add_run('Reduces perceived wait time and user anxiety')

# Improvement 2
doc.add_heading('DESIGN-02: Implement Keyboard Navigation for Style Presets', level=2)
p = doc.add_paragraph()
p.add_run('Current State: ').bold = True
p.add_run('Mouse-only selection')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Suggested Improvement: ').bold = True
p.add_run('Add arrow key navigation with visual focus indicator on selected preset')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Implementation: ').bold = True
p.add_run('Use roving tabindex pattern - one preset has tabindex="0", others have tabindex="-1", arrow keys move focus')

# Improvement 3
doc.add_heading('DESIGN-03: Add Estimated Generation Time', level=2)
p = doc.add_paragraph()
p.add_run('Suggested Improvement: ').bold = True
p.add_run('Show "Estimated time: ~1 minute" before starting generation')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Benefit: ').bold = True
p.add_run('Sets user expectations and allows them to decide if they want to proceed')

# Improvement 4
doc.add_heading('DESIGN-04: Add Format Support Indicator', level=2)
p = doc.add_paragraph()
p.add_run('Suggested Improvement: ').bold = True
p.add_run('When opening unsupported formats, show message: "AI Restyle is available for PNG and JPEG images"')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Benefit: ').bold = True
p.add_run('Reduces user confusion about missing features')

# Improvement 5
doc.add_heading('DESIGN-05: Add Keyboard Shortcuts', level=2)
p = doc.add_paragraph()
p.add_run('Suggested Shortcuts: ').bold = True
doc.add_paragraph('Ctrl+Z / Cmd+Z - Undo', style='List Bullet')
doc.add_paragraph('Ctrl+Y / Cmd+Y - Redo', style='List Bullet')
doc.add_paragraph('Ctrl+S / Cmd+S - Save copy', style='List Bullet')
doc.add_paragraph('Escape - Close/Back', style='List Bullet')

# Improvement 6
doc.add_heading('DESIGN-06: AI-Generated Alt Text for Gallery', level=2)
p = doc.add_paragraph()
p.add_run('Suggested Improvement: ').bold = True
p.add_run('Use existing AI capabilities to auto-generate descriptive alt text for all gallery images')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Benefit: ').bold = True
p.add_run('Major accessibility improvement for screen reader users')

# ============== SCREENSHOTS INDEX ==============
doc.add_heading('Screenshots Index', level=1)

screenshots_table = doc.add_table(rows=10, cols=2)
screenshots_table.style = 'Table Grid'
screenshots_data = [
    ('Filename', 'Description'),
    ('audit_01_keyboard_focus.png', 'Gallery keyboard focus indicator'),
    ('audit_02_image_viewer.png', 'Image viewer with full toolbar'),
    ('audit_03_restyle_panel.png', 'Restyle panel with style presets'),
    ('audit_04_keyboard_nav_issue.png', 'Keyboard navigation issue'),
    ('audit_05_generation_started.png', 'AI generation in progress'),
    ('audit_06_generation_complete.png', 'Generation complete - Pop Art style'),
    ('audit_07_after_undo.png', 'After Undo - original restored'),
    ('audit_08_after_redo.png', 'After Redo - styled image restored'),
    ('audit_09_unsaved_dialog.png', 'Unsaved changes confirmation dialog')
]
for i, row_data in enumerate(screenshots_data):
    for j, cell_text in enumerate(row_data):
        screenshots_table.rows[i].cells[j].text = cell_text
        if i == 0:
            screenshots_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

# ============== CONCLUSION ==============
doc.add_heading('Conclusion', level=1)

conclusion = doc.add_paragraph()
conclusion.add_run('The AI Restyle feature demonstrates solid core functionality with good UI design for the main workflows. ')
conclusion.add_run('However, several accessibility issues need attention to meet WCAG 2.1 Level A compliance:\n\n')
conclusion.add_run('Critical Issues:\n').bold = True
conclusion.add_run('1. Gallery images lacking alt text (48 images)\n')
conclusion.add_run('2. Style presets not keyboard navigable\n\n')
conclusion.add_run('Previously Known Bugs:\n').bold = True
conclusion.add_run('1. Stop button non-functional\n')
conclusion.add_run('2. WEBP format unsupported\n\n')
conclusion.add_run('Positive Findings:\n').bold = True
conclusion.add_run('- Latency is acceptable for AI operations\n')
conclusion.add_run('- Unsaved changes dialog is well-implemented\n')
conclusion.add_run('- Undo/Redo works correctly\n')
conclusion.add_run('- Style presets have proper alt text\n')
conclusion.add_run('- Focus management in dialogs is correct')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Report Generated: ').bold = True
p.add_run('January 27, 2026\n')
p.add_run('Automation Tool: ').bold = True
p.add_run('Claude Code with Playwright MCP\n')
p.add_run('Total Test Duration: ').bold = True
p.add_run('~8 minutes')

# Save document
output_path = str(Path(__file__).resolve().parent / 'Comprehensive_Bug_Bash_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
