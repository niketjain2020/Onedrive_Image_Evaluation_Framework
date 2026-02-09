from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import os

# Create document
doc = Document()

# Set up styles
title = doc.add_heading('AI Restyle Custom Prompt Testing Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Testing 12 candidate style prompts for potential product addition')
doc.add_paragraph(f'Date: January 27, 2026')
doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', level=1)
summary = doc.add_paragraph()
summary.add_run('Objective: ').bold = True
summary.add_run('Identify 3 new transformational style prompts to add to the AI Restyle feature alongside the existing 14 style presets.\n\n')
summary.add_run('Method: ').bold = True
summary.add_run('Tested 12 candidate prompts on sample images, capturing results and evaluating transformation quality.\n\n')
summary.add_run('Recommendation: ').bold = True
summary.add_run('Add Japanese Ukiyo-e, Surrealist Dreamscape, and Art Deco Poster styles.')

doc.add_page_break()

# Top 3 Recommendations
doc.add_heading('TOP 3 RECOMMENDED STYLES', level=1)

recommendations = [
    {
        'rank': '1',
        'name': 'Japanese Ukiyo-e Woodblock Print',
        'suggested_name': 'Ukiyo-e',
        'prompt': 'Japanese ukiyo-e woodblock print',
        'why': 'Unique cultural art style not covered by any existing preset. Produces authentic traditional Japanese aesthetics with stylized clouds and muted colors.',
        'screenshot': 'prompt05_japanese_ukiyoe.png'
    },
    {
        'rank': '2',
        'name': 'Surrealist Dreamscape',
        'suggested_name': 'Dreamscape',
        'prompt': 'Surrealist dreamscape',
        'why': 'Highest transformation level - creates entirely new magical environments with floating elements, multiple moons, and butterflies. Highly shareable results.',
        'screenshot': 'prompt08_surrealist_dreamscape.png'
    },
    {
        'rank': '3',
        'name': 'Art Deco Poster',
        'suggested_name': 'Art Deco',
        'prompt': 'Art Deco poster',
        'why': 'Creates poster-style artwork with geometric sunburst patterns, bold frames, and auto-generated text labels. Unique vintage aesthetic.',
        'screenshot': 'prompt09_art_deco_poster.png'
    }
]

screenshot_folder = str(Path(__file__).resolve().parent.parent.parent / '.playwright-mcp')

for rec in recommendations:
    doc.add_heading(f"#{rec['rank']}: {rec['name']}", level=2)

    p = doc.add_paragraph()
    p.add_run('Suggested Product Name: ').bold = True
    p.add_run(f'"{rec["suggested_name"]}"\n')
    p.add_run('Prompt Used: ').bold = True
    p.add_run(f'"{rec["prompt"]}"\n\n')
    p.add_run('Why This Style: ').bold = True
    p.add_run(rec['why'])

    # Add screenshot
    img_path = os.path.join(screenshot_folder, rec['screenshot'])
    if os.path.exists(img_path):
        doc.add_paragraph()
        doc.add_picture(img_path, width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

doc.add_page_break()

# All 12 Prompts Tested
doc.add_heading('ALL 12 PROMPTS TESTED', level=1)

all_prompts = [
    {
        'num': 1,
        'prompt': 'Renaissance oil painting',
        'quality': 'Excellent',
        'description': 'Classical art style with rich textures, dramatic lighting, and dark backgrounds reminiscent of Old Masters.',
        'screenshot': 'prompt01_renaissance_oil_painting.png'
    },
    {
        'num': 2,
        'prompt': 'Impressionist painting',
        'quality': 'Excellent',
        'description': 'Soft visible brushstrokes, warm color palette, and dreamy atmosphere typical of Monet/Renoir style.',
        'screenshot': 'prompt02_impressionist_painting.png'
    },
    {
        'num': 3,
        'prompt': 'Cinematic film still',
        'quality': 'Good',
        'description': 'Enhanced dramatic lighting with shallow depth of field. More photorealistic enhancement than artistic transformation.',
        'screenshot': 'prompt03_cinematic_film_still.png'
    },
    {
        'num': 4,
        'prompt': 'Wes Anderson aesthetic',
        'quality': 'Excellent',
        'description': 'Symmetrical composition, pastel color palette, and stylized backgrounds characteristic of Wes Anderson films.',
        'screenshot': 'prompt04_wes_anderson_aesthetic.png'
    },
    {
        'num': 5,
        'prompt': 'Japanese ukiyo-e woodblock print',
        'quality': 'Excellent',
        'description': 'Authentic traditional Japanese art style with stylized clouds, muted colors, and woodblock print textures.',
        'screenshot': 'prompt05_japanese_ukiyoe.png'
    },
    {
        'num': 6,
        'prompt': 'Chinese ink wash painting',
        'quality': 'Excellent',
        'description': 'Elegant monochromatic style with delicate brushwork and traditional red seal. Beautiful Eastern art aesthetic.',
        'screenshot': 'prompt06_chinese_ink_wash.png'
    },
    {
        'num': 7,
        'prompt': '1970s vintage film photography',
        'quality': 'Good',
        'description': 'Warm faded tones, film grain texture, and vintage color processing typical of 1970s photography.',
        'screenshot': 'prompt07_1970s_vintage_film.png'
    },
    {
        'num': 8,
        'prompt': 'Surrealist dreamscape',
        'quality': 'Excellent',
        'description': 'Magical transformation with floating elements, multiple moons, ethereal clouds, and butterflies. Very creative.',
        'screenshot': 'prompt08_surrealist_dreamscape.png'
    },
    {
        'num': 9,
        'prompt': 'Art Deco poster',
        'quality': 'Excellent',
        'description': 'Bold geometric sunburst patterns, decorative frame, and auto-generated "NATURE" text. 1920s-30s poster aesthetic.',
        'screenshot': 'prompt09_art_deco_poster.png'
    },
    {
        'num': 10,
        'prompt': 'Ethereal fantasy glowing light',
        'quality': 'Excellent',
        'description': 'Magical sparkles, glowing light effects, and dreamy purple atmosphere. Beautiful fantasy aesthetic.',
        'screenshot': 'prompt10_ethereal_fantasy.png'
    },
    {
        'num': 11,
        'prompt': 'Dark gothic fantasy',
        'quality': 'Excellent',
        'description': 'Moody dark atmosphere with dead trees, muted purple/gray tones, and dramatic gothic aesthetic.',
        'screenshot': 'prompt11_dark_gothic_fantasy.png'
    },
    {
        'num': 12,
        'prompt': 'Polaroid instant photo vintage',
        'quality': 'Good',
        'description': 'Classic Polaroid frame with white border and vintage washed-out color treatment.',
        'screenshot': 'prompt12_polaroid_vintage.png'
    }
]

for item in all_prompts:
    doc.add_heading(f"Prompt {item['num']}: \"{item['prompt']}\"", level=2)

    p = doc.add_paragraph()
    p.add_run('Quality Rating: ').bold = True
    p.add_run(f"{item['quality']}\n")
    p.add_run('Description: ').bold = True
    p.add_run(item['description'])

    # Add screenshot
    img_path = os.path.join(screenshot_folder, item['screenshot'])
    if os.path.exists(img_path):
        doc.add_paragraph()
        doc.add_picture(img_path, width=Inches(5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

doc.add_page_break()

# Gap Analysis
doc.add_heading('Gap Analysis vs. Existing Presets', level=1)

doc.add_paragraph('Current 12 presets in the product:')
existing = doc.add_paragraph()
existing.add_run('Movie Poster, Plush Toy, Anime, Graffiti, Crochet Art, Forest Scene, Cherry Blossoms, Neon Glow, Hologram, Doodle, Storybook, Glass Mosaic')

doc.add_paragraph()
doc.add_paragraph('Gaps filled by recommended additions:')

# Create table
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Category Gap'
header_cells[1].text = 'Current Coverage'
header_cells[2].text = 'Recommended Addition'

# Make headers bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
data = [
    ('Traditional Asian Art', 'None', 'Japanese Ukiyo-e'),
    ('Surreal/Fantasy Environments', 'None (only style filters)', 'Surrealist Dreamscape'),
    ('Vintage Graphic Design', 'Movie Poster (different style)', 'Art Deco Poster')
]

for i, row_data in enumerate(data):
    row = table.rows[i + 1]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]
    row.cells[2].text = row_data[2]

doc.add_paragraph()

# Conclusion
doc.add_heading('Conclusion', level=1)
conclusion = doc.add_paragraph()
conclusion.add_run('The three recommended styles (Japanese Ukiyo-e, Surrealist Dreamscape, and Art Deco Poster) provide:')
doc.add_paragraph('1. High transformation quality with distinct, shareable results', style='List Bullet')
doc.add_paragraph('2. Coverage of aesthetic categories not addressed by current presets', style='List Bullet')
doc.add_paragraph('3. Broad appeal across different user preferences', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Next steps: Test on portrait images to verify face handling and kid-image blocking behavior.')

# Save
output_path = str(Path(__file__).resolve().parent / 'Style_Prompt_Testing_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
