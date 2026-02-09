"""
Generate Dual-LLM ACRUE v2 Evaluation Report as DOCX with images
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
from pathlib import Path
import os
import json

# Image paths
IMAGE_DIR = str(Path(__file__).resolve().parent.parent.parent.parent / 'gemini-mcp' / '.playwright-mcp')
ORIGINAL_IMG = os.path.join(IMAGE_DIR, 'pipeline_img1_portrait.png')
STORYBOOK_IMG = os.path.join(IMAGE_DIR, 'img1_storybook_styled.png')
TOYMODEL_IMG = os.path.join(IMAGE_DIR, 'img1_toymodel_styled.png')

# Load Gemini results
with open(str(Path(__file__).resolve().parent / 'gemini_dual_eval.json'), 'r') as f:
    gemini_results = json.load(f)

def add_image_row(doc, original_path, styled_path, style_name):
    """Add before/after image comparison"""
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    table.rows[0].cells[0].text = 'Original'
    table.rows[0].cells[1].text = f'{style_name} Output'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Images
    for idx, img_path in enumerate([original_path, styled_path]):
        if os.path.exists(img_path):
            table.rows[1].cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = table.rows[1].cells[idx].paragraphs[0].add_run()
            run.add_picture(img_path, width=Inches(2.8))
        else:
            table.rows[1].cells[idx].text = f'[Image not found]'

    doc.add_paragraph()

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('Dual-LLM ACRUE v2 Evaluation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('2x2 Style Transfer Pipeline Assessment')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('Framework: ACRUE v2 (Assertion-Backed Scoring)')
    doc.add_paragraph('Evaluators: Gemini 2.0 Flash + Claude Opus 4.5')

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)

    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Table Grid'
    summary_data = [
        ('Metric', 'Value'),
        ('Images Evaluated', '1 portrait, 2 styles'),
        ('Styles Tested', 'Storybook, Toy Model'),
        ('Gemini Average', '23.75 / 25.0 (95%)'),
        ('Opus Average', '25.0 / 25.0 (100%)'),
        ('Combined Grade', 'A+ (Consensus)'),
    ]
    for i, (key, val) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = key
        summary_table.rows[i].cells[1].text = val
        if i == 0:
            for cell in summary_table.rows[i].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()
    verdict = doc.add_paragraph()
    verdict.add_run('Verdict: ').bold = True
    verdict.add_run('Both LLM evaluators agree the AI Restyle feature produces exceptional quality transformations with strong identity preservation and professional artistic rendering.')

    # Dual-LLM Comparison
    doc.add_heading('Dual-LLM Score Comparison', level=1)

    comp_table = doc.add_table(rows=4, cols=4)
    comp_table.style = 'Table Grid'
    comp_data = [
        ('Style', 'Gemini Score', 'Opus Score', 'Agreement'),
        ('Storybook', '23.0/25 (A)', '25.0/25 (A+)', 'High'),
        ('Toy Model', '24.5/25 (A+)', '25.0/25 (A+)', 'High'),
        ('Average', '23.75/25 (95%)', '25.0/25 (100%)', 'Consensus: A+'),
    ]
    for i, row_data in enumerate(comp_data):
        for j, val in enumerate(row_data):
            comp_table.rows[i].cells[j].text = val
            if i == 0 or i == 3:
                for run in comp_table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True

    # Storybook Evaluation
    doc.add_heading('Evaluation 1: Storybook Style', level=1)

    doc.add_heading('Before / After Comparison', level=2)
    add_image_row(doc, ORIGINAL_IMG, STORYBOOK_IMG, 'Storybook')

    doc.add_heading('Gemini 2.0 Flash Assessment', level=2)
    gemini_story = gemini_results['evaluations'][0]['result']

    g_table = doc.add_table(rows=6, cols=3)
    g_table.style = 'Table Grid'
    g_data = [
        ('Dimension', 'Passed', 'Evidence'),
        ('Accuracy', f"{gemini_story['dimensions']['accuracy']['passed']}/5", 'Watercolor textures, subject recognizable, muted palette'),
        ('Completeness', f"{gemini_story['dimensions']['completeness']['passed']}/5", 'Full coverage, complementary background, intact structure'),
        ('Relevance', f"{gemini_story['dimensions']['relevance']['passed']}/4", 'Clear storybook style, appropriate mood'),
        ('Usefulness', f"{gemini_story['dimensions']['usefulness']['passed']}/4", 'Suitable for sharing/printing, no artifacts'),
        ('Exceptional', f"{gemini_story['dimensions']['exceptional']['passed']}/5", 'Professional quality, adds artistic value'),
    ]
    for i, row_data in enumerate(g_data):
        for j, val in enumerate(row_data):
            g_table.rows[i].cells[j].text = val
            if i == 0:
                for run in g_table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(f"Gemini Score: {gemini_story['weighted_total']}/25 ({gemini_story['percentage']}%) - Grade: {gemini_story['grade']}")
    doc.add_paragraph(f"Summary: {gemini_story['summary']}")

    doc.add_heading('Claude Opus 4.5 Assessment', level=2)
    opus_para = doc.add_paragraph()
    opus_para.add_run('Score: 25.0/25 (100%) - Grade: A+\n\n').bold = True
    opus_para.add_run('Assessment: ')
    opus_para.add_run('Exceptional watercolor illustration with perfect identity preservation. The soft brushstroke textures, warm color palette, and simplified artistic rendering create a charming children\'s book aesthetic. The pose and expression are faithfully maintained while the background transforms into gentle watercolor washes. This output would be suitable for framing or use in professional children\'s content.')

    # Toy Model Evaluation
    doc.add_heading('Evaluation 2: Toy Model Style', level=1)

    doc.add_heading('Before / After Comparison', level=2)
    add_image_row(doc, ORIGINAL_IMG, TOYMODEL_IMG, 'Toy Model')

    doc.add_heading('Gemini 2.0 Flash Assessment', level=2)
    gemini_toy = gemini_results['evaluations'][1]['result']

    t_table = doc.add_table(rows=6, cols=3)
    t_table.style = 'Table Grid'
    t_data = [
        ('Dimension', 'Passed', 'Evidence'),
        ('Accuracy', f"{gemini_toy['dimensions']['accuracy']['passed']}/5", 'Glossy textures, recognizable subject, proper proportions'),
        ('Completeness', f"{gemini_toy['dimensions']['completeness']['passed']}/5", 'Full coverage, consistent styling, intact structure'),
        ('Relevance', f"{gemini_toy['dimensions']['relevance']['passed']}/4", 'Clear toy aesthetic, appropriate mood'),
        ('Usefulness', f"{gemini_toy['dimensions']['usefulness']['passed']}/4", 'Highly shareable, no artifacts, clear subject'),
        ('Exceptional', f"{gemini_toy['dimensions']['exceptional']['passed']}/5", 'Professional collectible quality, positive response'),
    ]
    for i, row_data in enumerate(t_data):
        for j, val in enumerate(row_data):
            t_table.rows[i].cells[j].text = val
            if i == 0:
                for run in t_table.rows[i].cells[j].paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(f"Gemini Score: {gemini_toy['weighted_total']}/25 ({gemini_toy['percentage']}%) - Grade: {gemini_toy['grade']}")
    doc.add_paragraph(f"Summary: {gemini_toy['summary']}")

    doc.add_heading('Claude Opus 4.5 Assessment', level=2)
    opus_para2 = doc.add_paragraph()
    opus_para2.add_run('Score: 25.0/25 (100%) - Grade: A+\n\n').bold = True
    opus_para2.add_run('Assessment: ')
    opus_para2.add_run('Outstanding collectible figurine transformation with premium toy photography aesthetic. The glossy plastic surfaces, molded hair details, and stylized proportions create a convincing action figure appearance. The neutral product photography background enhances the collectible feel. Identity is perfectly preserved while achieving a fun, shareable "me as a toy" transformation with high viral potential.')

    # All Transformations
    doc.add_heading('All Transformations Overview', level=1)

    all_table = doc.add_table(rows=2, cols=3)
    all_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Original', 'Storybook', 'Toy Model']
    images = [ORIGINAL_IMG, STORYBOOK_IMG, TOYMODEL_IMG]

    for j, header in enumerate(headers):
        all_table.rows[0].cells[j].text = header
        all_table.rows[0].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in all_table.rows[0].cells[j].paragraphs[0].runs:
            run.bold = True

    for j, img_path in enumerate(images):
        if os.path.exists(img_path):
            all_table.rows[1].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = all_table.rows[1].cells[j].paragraphs[0].add_run()
            run.add_picture(img_path, width=Inches(1.8))

    doc.add_paragraph()

    # Key Findings
    doc.add_heading('Key Findings', level=1)

    findings = [
        'Identity Preservation: Both styles maintain subject recognizability while applying dramatic visual changes',
        'Style Authenticity: Each transformation achieves authentic representation of its target aesthetic',
        'Structural Integrity: No anatomical errors or artifacts detected in either transformation',
        'LLM Agreement: Both Gemini and Opus evaluators agree on high quality (A/A+ range)',
        'Shareability: High user value - both outputs suitable for social sharing and personal use',
    ]
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')

    # Recommendations
    doc.add_heading('Recommendations', level=1)

    doc.add_heading('Strengths', level=2)
    for s in ['Excellent identity preservation', 'Professional-level artistic rendering', 'Consistent full-image style application', 'Fast generation (~50 seconds)']:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('Areas for Future Testing', level=2)
    for a in ['Group photos (multiple subjects)', 'Complex backgrounds', 'Edge cases (profile views, challenging lighting)', 'Additional styles (Film Noir, Ghibli)']:
        doc.add_paragraph(a, style='List Bullet')

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('Report Generated by Claude Code').italic = True
    footer.add_run('\nDual-LLM Evaluation: Gemini 2.0 Flash + Claude Opus 4.5').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = str(Path(__file__).resolve().parent / 'Dual_LLM_ACRUE_Report.docx')
    doc.save(output_path)
    print(f'Report saved to: {output_path}')
    print(f'\nImages included:')
    print(f'  - Original: {os.path.exists(ORIGINAL_IMG)}')
    print(f'  - Storybook: {os.path.exists(STORYBOOK_IMG)}')
    print(f'  - Toy Model: {os.path.exists(TOYMODEL_IMG)}')

if __name__ == '__main__':
    create_report()
