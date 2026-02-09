"""
Bug Bash Demo Report Generator
Generates a DOCX report for the AI Restyle Bug Bash Demo Session
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
from pathlib import Path
import os

def create_bugbash_report():
    doc = Document()

    # Title
    title = doc.add_heading('OneDrive Photos AI Restyle', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Bug Bash Demo Report', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date and metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}').italic = True
    meta.add_run('\n')
    meta.add_run('Testing Tool: Claude Code + Playwright MCP').italic = True

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run('This report documents a live bug bash demo session conducted using ')
    summary.add_run('Claude Code with Playwright MCP').bold = True
    summary.add_run(' for autonomous QA testing of the OneDrive Photos AI Restyle feature. ')
    summary.add_run('The demo successfully reproduced a known P1 bug and demonstrated end-to-end automated testing capabilities.')

    # Demo Flow
    doc.add_heading('Demo Flow', level=1)

    flow_table = doc.add_table(rows=7, cols=3)
    flow_table.style = 'Table Grid'

    # Header row
    header_cells = flow_table.rows[0].cells
    header_cells[0].text = 'Step'
    header_cells[1].text = 'Action'
    header_cells[2].text = 'Result'

    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        shading = cell._element.get_or_add_tcPr()

    # Data rows
    flow_data = [
        ('1', 'Navigate to OneDrive Photos', 'Gallery loaded successfully'),
        ('2', 'Open photo (IMG_3787.jpeg)', 'Photo viewer with toolbar appeared'),
        ('3', 'Click "Restyle with AI"', 'Restyle panel opened with 12 style presets'),
        ('4', 'Select "Anime" style', 'Prompt auto-populated with detailed instructions'),
        ('5', 'Click Send', 'AI generation started'),
        ('6', 'Click Stop button', 'BUG: Generation continued (not cancelled)'),
    ]

    for i, (step, action, result) in enumerate(flow_data, 1):
        row = flow_table.rows[i].cells
        row[0].text = step
        row[1].text = action
        row[2].text = result

    doc.add_paragraph()

    # Bug Found Section
    doc.add_heading('Bug Reproduced: Stop Button Non-Functional', level=1)

    bug_para = doc.add_paragraph()
    bug_para.add_run('Severity: ').bold = True
    bug_para.add_run('P1 - HIGH')

    bug_para2 = doc.add_paragraph()
    bug_para2.add_run('Category: ').bold = True
    bug_para2.add_run('Functional Bug')

    doc.add_heading('Description', level=2)
    desc = doc.add_paragraph(
        'The Stop button appears during AI generation but does not cancel the operation. '
        'When clicked, the generation continues to completion instead of being cancelled.'
    )

    doc.add_heading('Steps to Reproduce', level=2)
    steps = [
        'Navigate to https://onedrive.live.com/photos',
        'Open any supported image (PNG or JPEG)',
        'Click "Restyle with AI" in the toolbar',
        'Select any style preset (e.g., Anime)',
        'Click Send to start generation',
        'Click the Stop button while generation is in progress',
        'Observe: Generation continues instead of stopping'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_heading('Expected Behavior', level=2)
    doc.add_paragraph('Clicking Stop should immediately cancel the AI generation process.')

    doc.add_heading('Actual Behavior', level=2)
    doc.add_paragraph(
        'Stop button click has no effect. Status messages continue cycling through: '
        '"Pixels getting warmed up...", "Getting your photo ready...", '
        '"Mixing everything into one awesome combo...", "Good things take time...", '
        '"It\'s taking longer than expected..."'
    )

    doc.add_paragraph()

    # Style Presets Observed
    doc.add_heading('AI Restyle Style Presets Observed', level=1)

    presets = [
        'Movie Poster', 'Plush Toy', 'Anime', 'Graffiti',
        'Crochet Art', 'Forest Scene', 'Cherry Blossoms', 'Neon Glow',
        'Hologram', 'Doodle', 'Storybook', 'Glass Mosaic'
    ]

    preset_table = doc.add_table(rows=4, cols=4)
    preset_table.style = 'Table Grid'

    for i, preset in enumerate(presets):
        row = i // 4
        col = i % 4
        preset_table.rows[row].cells[col].text = preset

    doc.add_paragraph()

    # Screenshots
    doc.add_heading('Evidence Captured', level=1)

    evidence_table = doc.add_table(rows=2, cols=2)
    evidence_table.style = 'Table Grid'

    evidence_table.rows[0].cells[0].text = 'Screenshot'
    evidence_table.rows[0].cells[1].text = 'Description'
    evidence_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    evidence_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    evidence_table.rows[1].cells[0].text = 'demo_bugbash_anime_generating.png'
    evidence_table.rows[1].cells[1].text = 'Restyle panel with Anime style selected, showing generation in progress with Stop button visible'

    doc.add_paragraph()

    # Try to add screenshot if it exists
    screenshot_path = str(Path(__file__).resolve().parent.parent.parent / '.playwright-mcp' / 'demo_bugbash_anime_generating.png')
    if os.path.exists(screenshot_path):
        doc.add_heading('Screenshot: AI Generation in Progress', level=2)
        try:
            doc.add_picture(screenshot_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f'[Screenshot available at: {screenshot_path}]')

    doc.add_paragraph()

    # Performance Notes
    doc.add_heading('Performance Observations', level=1)

    perf_table = doc.add_table(rows=4, cols=2)
    perf_table.style = 'Table Grid'

    perf_data = [
        ('Metric', 'Value'),
        ('Gallery Load Time', '~1 second'),
        ('Restyle Panel Open', '<1 second'),
        ('AI Generation Time', '~50+ seconds (baseline), longer with network issues')
    ]

    for i, (metric, value) in enumerate(perf_data):
        perf_table.rows[i].cells[0].text = metric
        perf_table.rows[i].cells[1].text = value
        if i == 0:
            perf_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            perf_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Testing Tool Benefits
    doc.add_heading('Claude Code + Playwright MCP Capabilities Demonstrated', level=1)

    capabilities = [
        'Autonomous Navigation - Navigated to OneDrive Photos without manual guidance',
        'UI Interaction - Clicked buttons, selected style presets, triggered AI generation',
        'Bug Detection - Identified and reproduced P1 Stop Button bug',
        'Evidence Collection - Captured screenshots automatically',
        'Accessibility Snapshot - Read page structure via accessibility tree',
        'Real-time Monitoring - Tracked status message changes during generation',
        'Report Generation - Created this DOCX report programmatically'
    ]

    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')

    doc.add_paragraph()

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion = doc.add_paragraph()
    conclusion.add_run(
        'This bug bash demo successfully demonstrated the power of autonomous QA testing '
        'using Claude Code with Playwright MCP. The system was able to navigate the OneDrive '
        'Photos interface, test the AI Restyle feature, reproduce a known P1 bug, and '
        'capture evidence - all without manual intervention. This approach enables faster, '
        'more thorough testing coverage and can be integrated into CI/CD pipelines for '
        'continuous quality monitoring.'
    )

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Generated by Claude Code + Playwright MCP').italic = True
    footer.add_run('\n')
    footer.add_run('Autonomous QA Testing Demo').italic = True

    # Save the document
    output_path = str(Path(__file__).resolve().parent / 'BugBash_Demo_Report.docx')
    doc.save(output_path)
    print(f'Report saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_bugbash_report()
