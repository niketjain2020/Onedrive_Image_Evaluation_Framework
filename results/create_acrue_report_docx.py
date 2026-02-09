"""
Generate ACRUE v2 Evaluation Report as DOCX with images
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
from pathlib import Path
import os

# Image paths
IMAGE_DIR = str(Path(__file__).resolve().parent.parent.parent / '.playwright-mcp')
ORIGINAL_IMG = os.path.join(IMAGE_DIR, 'qa_step1_viewer_opened.png')
STORYBOOK_IMG = os.path.join(IMAGE_DIR, 'img1-12-storybook.png')
ANIME_IMG = os.path.join(IMAGE_DIR, 'img1-03-anime.png')

def add_image_comparison(doc, original_path, restyled_path, style_name):
    """Add side-by-side image comparison table"""
    # Create a 1x2 table for images
    img_table = doc.add_table(rows=2, cols=2)
    img_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    img_table.rows[0].cells[0].text = 'Original'
    img_table.rows[0].cells[1].text = f'{style_name} Output'
    for cell in img_table.rows[0].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True

    # Add images
    if os.path.exists(original_path):
        img_table.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_table.rows[1].cells[0].paragraphs[0].add_run()
        run.add_picture(original_path, width=Inches(2.8))
    else:
        img_table.rows[1].cells[0].text = f'[Image not found: {original_path}]'

    if os.path.exists(restyled_path):
        img_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_table.rows[1].cells[1].paragraphs[0].add_run()
        run.add_picture(restyled_path, width=Inches(2.8))
    else:
        img_table.rows[1].cells[1].text = f'[Image not found: {restyled_path}]'

    doc.add_paragraph()

def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('ACRUE v2 Evaluation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('OneDrive Photos AI Restyle Quality Assessment')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d")}')
    doc.add_paragraph('Framework: ACRUE v2 (Assertion-Backed Scoring)')
    doc.add_paragraph('Evaluator: Gemini 2.0 Flash via MCP')

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)

    summary_table = doc.add_table(rows=7, cols=2)
    summary_table.style = 'Table Grid'
    summary_data = [
        ('Metric', 'Value'),
        ('Images Evaluated', '2'),
        ('Styles Tested', 'Storybook, Anime'),
        ('Average Score', '25.0 / 25.0 (100%)'),
        ('Average Grade', 'A+'),
        ('Total Assertions', '46'),
        ('Assertions Passed', '46 (100%)'),
    ]
    for i, (key, val) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = key
        summary_table.rows[i].cells[1].text = val
        if i == 0:
            for cell in summary_table.rows[i].cells:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    verdict = doc.add_paragraph()
    verdict.add_run('Verdict: ').bold = True
    verdict.add_run('The OneDrive Photos AI Restyle feature demonstrates exceptional quality across tested styles, with perfect assertion pass rates and maximum scores in all ACRUE dimensions.')

    # Framework Section
    doc.add_heading('Evaluation Framework', level=1)
    doc.add_heading('ACRUE v2 Dimensions & Weights', level=2)

    dim_table = doc.add_table(rows=7, cols=4)
    dim_table.style = 'Table Grid'
    dim_data = [
        ('Dimension', 'Weight', 'Max Score', 'Description'),
        ('A - Accuracy', '1.0', '5.0', 'Identity preservation + style authenticity'),
        ('C - Completeness', '1.0', '5.0', 'Full coverage + structural integrity'),
        ('R - Relevance', '0.5', '2.5', 'Style match + mood alignment'),
        ('U - Usefulness', '0.5', '2.5', 'Share-ready + artifact-free'),
        ('E - Exceptional', '2.0', '10.0', 'Professional quality + wow factor'),
        ('TOTAL', '5.0', '25.0', ''),
    ]
    for i, row_data in enumerate(dim_data):
        for j, val in enumerate(row_data):
            dim_table.rows[i].cells[j].text = val
            if i == 0 or i == 6:
                dim_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    # Storybook Evaluation
    doc.add_heading('Evaluation 1: Storybook Style', level=1)

    # Add image comparison for Storybook
    doc.add_heading('Before / After Comparison', level=2)
    add_image_comparison(doc, ORIGINAL_IMG, STORYBOOK_IMG, 'Storybook')

    doc.add_heading('Assertion Results Summary', level=2)

    story_assertions = [
        ('A - Accuracy', '5/5', 'Watercolor textures, identity preserved, warm palette, expressions intact, professional rendering'),
        ('C - Completeness', '5/5', 'Full coverage, cohesive background, structural integrity, balanced saturation, consistent style'),
        ('R - Relevance', '4/4', 'Resembles children\'s books, gentle style, warm mood, age-appropriate'),
        ('U - Usefulness', '4/4', 'Suitable for nursery decor, artifact-free, family-friendly, print-ready'),
        ('E - Exceptional', '5/5', 'Professional artist quality, evokes warmth/joy, highly shareable, frame-worthy, standout quality'),
    ]

    story_table = doc.add_table(rows=6, cols=3)
    story_table.style = 'Table Grid'
    story_table.rows[0].cells[0].text = 'Dimension'
    story_table.rows[0].cells[1].text = 'Passed'
    story_table.rows[0].cells[2].text = 'Key Evidence'
    for cell in story_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    for i, (dim, passed, evidence) in enumerate(story_assertions):
        story_table.rows[i+1].cells[0].text = dim
        story_table.rows[i+1].cells[1].text = passed
        story_table.rows[i+1].cells[2].text = evidence

    doc.add_paragraph()
    score_para = doc.add_paragraph()
    score_para.add_run('Score: 25.0 / 25.0 (100%) - Grade: A+').bold = True

    # Anime Evaluation
    doc.add_heading('Evaluation 2: Anime Style', level=1)

    # Add image comparison for Anime
    doc.add_heading('Before / After Comparison', level=2)
    add_image_comparison(doc, ORIGINAL_IMG, ANIME_IMG, 'Anime')

    doc.add_heading('Assertion Results Summary', level=2)

    anime_assertions = [
        ('A - Accuracy', '5/5', 'Subject recognizable, authentic anime style, pose preserved, tasteful elements, professional quality'),
        ('C - Completeness', '5/5', 'Full style coverage, complementary background, structural integrity, balanced saturation, consistent rendering'),
        ('R - Relevance', '4/4', 'Clear anime style, cheerful tone, expected mood, no inconsistencies'),
        ('U - Usefulness', '4/4', 'Social media ready, artifact-free, subject visible, adequate resolution'),
        ('E - Exceptional', '5/5', 'Professional quality, adds artistic value, highly shareable, standout quality, positive emotional response'),
    ]

    anime_table = doc.add_table(rows=6, cols=3)
    anime_table.style = 'Table Grid'
    anime_table.rows[0].cells[0].text = 'Dimension'
    anime_table.rows[0].cells[1].text = 'Passed'
    anime_table.rows[0].cells[2].text = 'Key Evidence'
    for cell in anime_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    for i, (dim, passed, evidence) in enumerate(anime_assertions):
        anime_table.rows[i+1].cells[0].text = dim
        anime_table.rows[i+1].cells[1].text = passed
        anime_table.rows[i+1].cells[2].text = evidence

    doc.add_paragraph()
    score_para2 = doc.add_paragraph()
    score_para2.add_run('Score: 25.0 / 25.0 (100%) - Grade: A+').bold = True

    # Comparative Analysis
    doc.add_heading('Comparative Analysis', level=1)

    # All images side by side
    doc.add_heading('All Transformations', level=2)
    all_img_table = doc.add_table(rows=2, cols=3)
    all_img_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Original', 'Storybook', 'Anime']
    images = [ORIGINAL_IMG, STORYBOOK_IMG, ANIME_IMG]

    for j, header in enumerate(headers):
        all_img_table.rows[0].cells[j].text = header
        all_img_table.rows[0].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        all_img_table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

    for j, img_path in enumerate(images):
        if os.path.exists(img_path):
            all_img_table.rows[1].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = all_img_table.rows[1].cells[j].paragraphs[0].add_run()
            run.add_picture(img_path, width=Inches(1.9))

    doc.add_paragraph()

    comp_table = doc.add_table(rows=3, cols=3)
    comp_table.style = 'Table Grid'
    comp_data = [
        ('Metric', 'Storybook', 'Anime'),
        ('Pass Rate', '23/23 (100%)', '23/23 (100%)'),
        ('Grade', 'A+', 'A+'),
    ]
    for i, row_data in enumerate(comp_data):
        for j, val in enumerate(row_data):
            comp_table.rows[i].cells[j].text = val
            if i == 0:
                comp_table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

    # Key Findings
    doc.add_heading('Key Findings', level=1)

    findings = [
        'Identity Preservation: Both styles successfully preserve subject recognition while applying dramatic visual transformations',
        'Style Authenticity: Each style achieves authentic representation of its target aesthetic',
        'Structural Integrity: No broken limbs, warped faces, or anatomical errors detected',
        'Professional Quality: Both outputs meet professional-grade standards',
        'Shareability: High user value - images suitable for social sharing and printing',
    ]

    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')

    # Recommendations
    doc.add_heading('Recommendations', level=1)

    doc.add_heading('Strengths to Maintain', level=2)
    strengths = [
        'Excellent identity preservation across styles',
        'Consistent full-image style application',
        'High-quality structural integrity',
        'Professional-level artistic rendering',
    ]
    for s in strengths:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('Areas for Future Testing', level=2)
    future = [
        'Test with group photos (multiple subjects)',
        'Test with complex backgrounds',
        'Test with challenging lighting conditions',
        'Evaluate additional styles (Cyberpunk, Oil Painting, etc.)',
    ]
    for f in future:
        doc.add_paragraph(f, style='List Bullet')

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('Report Generated by Claude Code + Gemini MCP').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = str(Path(__file__).resolve().parent / 'ACRUE_v2_Evaluation_Report.docx')
    doc.save(output_path)
    print(f'Report saved to: {output_path}')
    print(f'Images included:')
    print(f'  - Original: {os.path.exists(ORIGINAL_IMG)}')
    print(f'  - Storybook: {os.path.exists(STORYBOOK_IMG)}')
    print(f'  - Anime: {os.path.exists(ANIME_IMG)}')

if __name__ == '__main__':
    create_report()
